# receipt_app.py
import re
import database
import mysql.connector
from flask import (
    Flask,
    render_template,
    request,
    redirect,
    url_for,
    abort,
    make_response,
    jsonify,
    session,
    flash,
    send_file,
    g,
)
from datetime import datetime
import os
import io
import io
import base64
import smtplib
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
import json
from pathlib import Path
from urllib.parse import urlparse, quote as urlquote
import requests
from provision_tenant import provision_new_tenant
import openpyxl
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from werkzeug.security import generate_password_hash, check_password_hash

# Playwright will be imported lazily; if missing, we show a helpful message at runtime.
try:
    from playwright.sync_api import sync_playwright
except Exception:
    sync_playwright = None

app = Flask(__name__)
app.config["UPLOAD_FOLDER"] = "static/images"
# Secret key for sessions + flash messages
app.secret_key = "CHANGE_THIS_SECRET_KEY_123456789"  # Change before deployment


# -------------------------------
# Database initialization
# -------------------------------
def init_db():
    conn = database.get_db_connection()
    c = conn.cursor()
    c.execute(
        """
    CREATE TABLE IF NOT EXISTS receipts (
        id INT AUTO_INCREMENT PRIMARY KEY,
        no VARCHAR(255),
        project_name VARCHAR(255),
        date VARCHAR(255),
        venture VARCHAR(255),
        customer_name VARCHAR(255),
        amount_numeric DOUBLE,
        amount_words TEXT,
        plot_no VARCHAR(255),
        square_yards VARCHAR(255),
        purpose TEXT,
        drawn_bank VARCHAR(255),
        branch VARCHAR(255),
        payment_mode VARCHAR(255),
        instrument_no VARCHAR(255),
        created_at VARCHAR(255)
    );
    """
    )
    conn.commit()
    conn.close()
    
    # Run migrations
    migrate_commissions_table()
    migrate_users_table()
    init_users()
    init_pending_receipts()


def migrate_commissions_table():
    """Add name columns and sq_yards to commissions table if they don't exist"""
    conn = database.get_db_connection()
    c = conn.cursor()
    
    # Get existing columns
    c.execute("SHOW COLUMNS FROM commissions")
    existing_columns = [row['Field'] for row in database.fetch_all(c)]
    
    # Add missing columns
    if 'cgm_name' not in existing_columns:
        c.execute("ALTER TABLE commissions ADD COLUMN cgm_name TEXT")
    if 'srgm_name' not in existing_columns:
        c.execute("ALTER TABLE commissions ADD COLUMN srgm_name TEXT")
    if 'gm_name' not in existing_columns:
        c.execute("ALTER TABLE commissions ADD COLUMN gm_name TEXT")

    if 'sq_yards' not in existing_columns:
        c.execute("ALTER TABLE commissions ADD COLUMN sq_yards REAL")
    if 'project_name' not in existing_columns:
        c.execute("ALTER TABLE commissions ADD COLUMN project_name TEXT")
    if 'commission_breakdown' not in existing_columns:
        c.execute("ALTER TABLE commissions ADD COLUMN commission_breakdown TEXT")
    
    conn.commit()
    conn.close()


def migrate_users_table():
    """Add permission columns to users table if they don't exist"""
    conn = database.get_db_connection()
    c = conn.cursor()
    
    # Check if users table exists
    c.execute("SELECT table_name FROM information_schema.tables WHERE table_schema = DATABASE() AND table_name = 'users'")
    if not database.fetch_one(c):
        conn.close()
        return
    
    # Get existing columns
    c.execute("SHOW COLUMNS FROM users")
    existing_columns = [row['Field'] for row in database.fetch_all(c)]
    
    # Add permission columns if they don't exist
    if 'can_search_receipts' not in existing_columns:
        try:
            c.execute("ALTER TABLE users ADD COLUMN can_search_receipts INTEGER NOT NULL DEFAULT 0")
            print("Added column can_search_receipts to users table")
        except database.OperationalError as e:
            print(f"Column can_search_receipts might already exist: {e}")
    
    conn.commit()
    conn.close()



# -------------------------------
# Helpers
# -------------------------------
def init_users():
    """Create users table and ensure at least one admin user exists.

    Columns:
      - username: unique login name
      - password_hash: salted hash created with werkzeug.security
      - role: 'admin' or 'user'
      - can_view_dashboard: 0/1 flag
    """
    conn = database.get_db_connection()
    c = conn.cursor()
    c.execute(
        """
        CREATE TABLE IF NOT EXISTS users (
            id INT AUTO_INCREMENT PRIMARY KEY,
            username VARCHAR(255) UNIQUE NOT NULL,
            password_hash VARCHAR(255) NOT NULL,
            role VARCHAR(50) NOT NULL DEFAULT 'user',
            can_view_dashboard TINYINT(1) NOT NULL DEFAULT 0,
            created_at VARCHAR(255) NOT NULL
        );
        """
    )
    conn.commit()

    # Ensure at least one admin user exists
    c.execute("SELECT COUNT(*) FROM users WHERE role = 'admin'")
    count_admin = database.fetch_one(c)[0] or 0
    if count_admin == 0:
        # Seed default admin user; you should change this password after first login.
        created_at = datetime.utcnow().isoformat()
        default_username = "admin"
        default_password = "password123"
        # Use pbkdf2:sha256 explicitly to avoid relying on hashlib.scrypt support
        pwd_hash = generate_password_hash(default_password, method="pbkdf2:sha256")
        try:
            c.execute(
                "INSERT INTO users (username, password_hash, role, can_view_dashboard, created_at) VALUES (%s, %s, 'admin', 1, %s)",
                (default_username, pwd_hash, created_at),
            )
            conn.commit()
        except database.IntegrityError:
            # Another row with same username was just created; ignore
            pass

    conn.close()


def init_pending_receipts():
    """Create pending_receipts table for approval workflow.
    
    This table stores receipts submitted by non-admin users that need admin approval.
    Once approved, the receipt is moved to the main receipts table.
    """
    conn = database.get_db_connection()
    c = conn.cursor()
    c.execute(
        """
        CREATE TABLE IF NOT EXISTS pending_receipts (
            id INT AUTO_INCREMENT PRIMARY KEY,
            no VARCHAR(255),
            project_name VARCHAR(255),
            date VARCHAR(255),
            venture VARCHAR(255),
            customer_name VARCHAR(255),
            amount_numeric DOUBLE,
            amount_words TEXT,
            plot_no VARCHAR(255),
            square_yards VARCHAR(255),
            purpose TEXT,
            drawn_bank VARCHAR(255),
            branch VARCHAR(255),
            payment_mode VARCHAR(255),
            instrument_no VARCHAR(255),
            submitted_by VARCHAR(255),
            submitted_at VARCHAR(255),
            status VARCHAR(50) DEFAULT 'pending',
            admin_notes TEXT
        );
        """
    )
    conn.commit()
    conn.close()


def format_inr(number):
    try:
        n = int(round(float(number)))
    except:
        return str(number or "")
    s = str(n)
    if len(s) <= 3:
        return s
    last3 = s[-3:]
    rem = s[:-3]
    parts = []
    while len(rem) > 2:
        parts.append(rem[-2:])
        rem = rem[:-2]
    if rem:
        parts.append(rem)
    parts.reverse()
    return ",".join(parts) + "," + last3


def number_to_words(n):
    try:
        from num2words import num2words

        out = num2words(int(float(n)), lang="en_IN")
        return out.replace("  ", " ").strip().title() + " Only"
    except Exception:
        return f"{format_inr(n)} Only"

def format_currency(amount):
    return f"Rs. {amount:,.2f}"


def dict_from_row(row):
    if row is None:
        return None
    return dict(row)


def _safe_filename_fragment(s: str, fallback: str):
    """
    Produce a safe filename fragment from `s`. Returns `fallback` if result is empty.
    Keeps letters, numbers, dot, underscore and dash. Replaces sequences of other chars with '_'.
    """
    if not s:
        return fallback
    frag = re.sub(r'[^A-Za-z0-9._-]+', '_', str(s)).strip('_')
    return frag or fallback


def get_column_index(table_name: str, column_name: str):
    """Return the 0-based column index for column_name in table_name, or None if not found/errors."""
    try:
        conn = database.get_db_connection()
        c = conn.cursor()
        c.execute(f"SHOW COLUMNS FROM {table_name}")
        rows = database.fetch_all(c)
        for i, row in enumerate(rows):
            # MySQL SHOW COLUMNS returns Field as first column
            # Using dict access from MySQLRow
            if row['Field'] == column_name:
                return i
        conn.close()
    except Exception:
        pass
    return None


# -------------------------------
# Project list helper
# -------------------------------
def get_projects():
    """
    Return a sorted list of project names from the 'projects' table.
    This function tries two common column names: 'name' and 'project_name'.
    If neither works it returns an empty list (safe fallback).
    """
    try:
        conn = database.get_db_connection()
        c = conn.cursor()
        
        # First, let's check if the projects table exists and has data
        c.execute("SHOW TABLES LIKE 'projects'")
        table_exists = c.fetchone()
        print(f"Projects table exists: {table_exists is not None}")
        
        if table_exists:
            c.execute("SELECT COUNT(*) FROM projects")
            count_result = c.fetchone()
            print(f"Projects count: {count_result}")
            
            # Show sample data
            c.execute("SELECT * FROM projects LIMIT 3")
            sample_data = c.fetchall()
            print(f"Sample data: {sample_data}")

        # Try 'name' first (MySQL doesn't support COLLATE NOCASE, use ORDER BY name)
        try:
            c.execute("SELECT name FROM projects WHERE name IS NOT NULL ORDER BY name")
            rows = database.fetch_all(c)
            print(f"Rows with 'name' column: {rows}")
            if rows:
                conn.close()
                # MySQLRow is dict-like, access by column name
                return [str(r['name']).strip() for r in rows if r and r.get('name') and str(r['name']).strip()]
        except Exception as e:
            print(f"Error with 'name' column: {e}")
            pass

        # Try 'project_name' column
        try:
            c.execute("SELECT project_name FROM projects WHERE project_name IS NOT NULL ORDER BY project_name")
            rows = database.fetch_all(c)
            print(f"Rows with 'project_name' column: {rows}")
            if rows:
                conn.close()
                return [str(r['project_name']).strip() for r in rows if r and r.get('project_name') and str(r['project_name']).strip()]
        except Exception as e:
            print(f"Error with 'project_name' column: {e}")
            pass

        conn.close()
        return []

    except Exception as e:
        print(f"Error in get_projects(): {e}")
        return []


# -------------------------------
# LOGIN SYSTEM
# -------------------------------


def _is_local_path(target: str):
    """
    Very small safeguard: ensure redirect target is local (path-only).
    Accepts paths like '/foo' or relative 'receipt/1'; rejects full external URLs.
    """
    if not target:
        return False
    parsed = urlparse(target)
    # parsed.netloc empty means not an absolute external URL
    return parsed.netloc == ""


@app.route("/login", methods=["GET", "POST"])
def login():
    """
    Renders login page on GET.
    Validates username/password on POST.
    Redirects to the next URL or home.
    """
    next_url = request.form.get("next") or request.args.get("next") or url_for("index")

    if request.method == "POST":
        username = request.form.get("username", "").strip()
        password = request.form.get("password", "").strip()

        # Look up user in DB
        conn = database.get_db_connection()
        c = conn.cursor()
        print(f"DEBUG: Attempting login for user: {username}")
        c.execute("SELECT id, username, password_hash, role, can_view_dashboard, can_search_receipts, can_view_vishvam_layout FROM users WHERE username = %s", (username,))
        row = database.fetch_one(c)
        conn.close()

        if row:
            # Fix: Access by key instead of unpacking (which yields keys for dict-like objects)
            user_id = row['id']
            uname = row['username']
            pwd_hash = row['password_hash']
            role = row['role']
            can_view_dash = row['can_view_dashboard']
            can_search = row['can_search_receipts']
            can_view_vishvam = row.get('can_view_vishvam_layout', 0)
            
            print(f"DEBUG: User found. ID: {user_id}, Role: {role}")
            # print(f"DEBUG: Stored hash: {pwd_hash}")
            try:
                valid = check_password_hash(pwd_hash, password)
                print(f"DEBUG: Password verification result: {valid}")
            except Exception as e:
                print(f"DEBUG: Password verification error: {e}")
                valid = False
            
            if valid:
                session.clear()
                session["logged_in"] = True
                session["user_id"] = user_id
                session["username"] = uname
                session["user"] = uname  # For navbar display
                session["role"] = role or "user"
                session["can_view_dashboard"] = bool(can_view_dash)
                session["can_search_receipts"] = bool(can_search)
                session["can_view_vishvam_layout"] = bool(can_view_vishvam)

                # Only redirect to local paths for safety
                if _is_local_path(next_url):
                    # If admin and next is index (default), redirect to dashboard
                    if role == 'admin' and (next_url == url_for("index") or next_url == "/"):
                        return redirect(url_for("dashboard"))
                    return redirect(next_url)
                
                if role == 'admin':
                    return redirect(url_for("dashboard"))
                return redirect(url_for("index"))
        else:
            print("DEBUG: User not found in database")

        flash("Invalid username or password", "danger")
        return render_template("login.html", next=next_url)

    return render_template("login.html", next=next_url)


@app.route("/logout")
def logout():
    session.clear()
    return redirect(url_for("login"))


@app.before_request
def tenant_routing():
    """
    SaaS Middleware:
    1. Determine Subdomain (e.g. 'srinidhi' from 'srinidhi.plotpro.in')
    2. Query Master DB for Tenant Config
    3. Configure Flask Global (g) for subsequent DB calls
    """
    # SKIP for static files to save DB hits
    if request.path.startswith('/static'):
        return

    host = request.headers.get("Host", "").split(":")[0].lower()  # remove port if any, lowercase
    
    # Extract subdomain
    # Logic: if host is 'srinidhi.plotpro.in', subdomain is 'srinidhi'
    # If host is 'localhost' or '127.0.0.1', treat as 'localhost' tenant
    
    if host in ["localhost", "127.0.0.1"]:
        subdomain = "localhost"
    elif host.count(".") >= 2:
        subdomain = host.split(".")[0]
    else:
        # Fallback or main domain (e.g. plotpro.in)
        subdomain = "www"
    
    import flask
    print(f"DEBUG: Host={host}, Count={host.count('.')}, Subdomain={subdomain}") # DEBUG LOG
    flask.g.subdomain = subdomain

    # Connect to Master DB
    try:
        # Connect using ENV vars which point to Master/Default DB
        master_conn = mysql.connector.connect(
            host=os.getenv("DB_HOST", "localhost"),
            user=os.getenv("DB_USER", "root"),
            password=os.getenv("DB_PASSWORD", ""),
            database="plotpro_master" # Explicitly connect to Master
        )
        c = master_conn.cursor(dictionary=True)
        
        c.execute("SELECT * FROM tenants WHERE subdomain = %s", (subdomain,))
        tenant = c.fetchone()
        master_conn.close()
        
        if tenant:
            # Store tenant config in g
            import flask
            flask.g.tenant = tenant
            flask.g.tenant_db_config = {
                "host": tenant["db_host"],
                "user": tenant["db_user"],
                "password": tenant["db_password"],
                "database": tenant["db_name"]
            }
        else:
            # Tenant not found. 
            # In production, redirect to 404 or Signup.
            # For dev simplicity, if subdomain='localhost' failed, we warn.
            if subdomain == "localhost":
               print("WARNING: 'localhost' tenant not found in Master DB. Please run init_saas_master.py")
            pass
            
    except Exception as e:
        print(f"SaaS Routing Error: {e}")
        # Fail gracefully? Or crash?
        # For now, proceed. If get_db_connection is called without config, it falls back to ENV.


@app.before_request
def enforce_login():
    """
    Protects all routes except /login and static file access and a few public endpoints.
    """
    # endpoints that should be accessible without login
    public = {"login", "static", "plot_lookup", "api.plot_lookup", "_routes", "list_routes", "contact_lead"}
    
    # Allow Landing Page on root domain (www) without login
    if getattr(g, 'subdomain', 'www') == 'www' and request.endpoint == 'index':
         return
    
    # Allow Super Admin routes (they have their own auth)
    if request.path.startswith('/superadmin'):
        return

    # in some cases request.endpoint can be None (favicon or malformed); allow login then
    if request.endpoint is None:
        return None
    if request.endpoint in public:
        return None
    if not session.get("logged_in"):
        return redirect(url_for("login", next=request.url))

    # Protect admin-only views (user management and future admin tools)
    admin_only = {"user_management"}
    if request.endpoint in admin_only and session.get("role") != "admin":
        abort(403)


# -------------------------------
# Routes
# -------------------------------
@app.route("/")
def index():
    # SaaS: If on main domain (www), show Landing Page
    if getattr(g, 'subdomain', 'www') == 'www':
        return render_template("landing_page.html")

    # App: Ensure login (since we bypassed enforce_login for www)
    if "user" not in session:
        return redirect(url_for("login"))

    # If user only has Vishvam access (not dashboard), redirect them to Vishvam Layout
    if not (session.get("role") == "admin" or session.get("can_view_dashboard")):
        if session.get("can_view_vishvam_layout"):
            return redirect(url_for("plot_layout_viewer", project_name="Vishvam"))
    
    conn = database.get_db_connection()
    c = conn.cursor()
    c.execute("SELECT id, no, customer_name, date FROM receipts ORDER BY id DESC LIMIT 3")
    rows = database.fetch_all(c)
    conn.close()
    projects = get_projects()
    return render_template("form.html", recent=rows, r=None, edit_mode=False, projects=projects)

@app.route("/api/contact_lead", methods=["POST"])
def contact_lead():
    """
    Handle contact form submission from Landing Page.
    1. Save to Master DB (leads table)
    2. Send Email (if configured)
    """
    data = request.json
    if not data:
        return jsonify({"error": "No data provided"}), 400
        
    try:
        # 1. Save to Database
        master_conn = mysql.connector.connect(
            host=os.getenv("DB_HOST", "localhost"),
            user=os.getenv("DB_USER", "root"),
            password=os.getenv("DB_PASSWORD", ""),
            database="plotpro_master"
        )
        c = master_conn.cursor()
        c.execute("""
            INSERT INTO leads (first_name, last_name, email, phone, company, plot_count, message)
            VALUES (%s, %s, %s, %s, %s, %s, %s)
        """, (
            data.get("firstName"), 
            data.get("lastName"), 
            data.get("email"), 
            data.get("phone"), 
            data.get("company"), 
            data.get("plotCount"), 
            data.get("message")
        ))
        master_conn.commit()
        master_conn.close()
        
        # 2. Send Email (Best Effort)
        smtp_server = os.getenv("MAIL_SERVER")
        smtp_port = os.getenv("MAIL_PORT")
        smtp_user = os.getenv("MAIL_USERNAME")
        smtp_pass = os.getenv("MAIL_PASSWORD")
        
        if smtp_server and smtp_user and smtp_pass:
            try:
                msg = MIMEMultipart()
                sender = smtp_user
                recipient = "sales@plotpro.in"
                msg['From'] = f"PlotPro Website <{sender}>"
                msg['To'] = recipient
                msg['Subject'] = f"New Lead: {data.get('firstName')} ({data.get('company')})"
                
                body = f"""
                New Demo Request Verification:
                ---------------------------
                Name: {data.get('firstName')} {data.get('lastName')}
                Company: {data.get('company')}
                Phone: {data.get('phone')}
                Email: {data.get('email')}
                Plots: {data.get('plotCount')}
                
                Message:
                {data.get('message')}
                """
                msg.attach(MIMEText(body, 'plain'))
                
                server = smtplib.SMTP(smtp_server, int(smtp_port) if smtp_port else 587)
                server.starttls()
                server.login(smtp_user, smtp_pass)
                server.sendmail(sender, recipient, msg.as_string())
                server.quit()
                print("Lead notification email sent.")
            except Exception as e:
                print(f"Failed to send email: {e}")
        
        return jsonify({"message": "Success", "status": "saved"}), 200

    except Exception as e:
        print(f"Lead API Error: {e}")
        return jsonify({"error": str(e)}), 500

# --- Super Admin Routes ---

@app.route("/superadmin/login", methods=["GET", "POST"])
def superadmin_login():
    if request.method == "POST":
        password = request.form.get("password")
        # Simple env-based auth
        correct_pass = os.getenv("SUPER_ADMIN_PASSWORD", "plotpro_super_2024")
        if password == correct_pass:
            session["is_super_admin"] = True
            return redirect(url_for("superadmin_dashboard"))
        else:
            flash("Invalid Password", "error")
    
    return render_template("superadmin_login.html")

@app.route("/superadmin/logout")
def superadmin_logout():
    session.pop("is_super_admin", None)
    return redirect(url_for("superadmin_login"))

@app.route("/superadmin/dashboard", methods=["GET", "POST"])
def superadmin_dashboard():
    if not session.get("is_super_admin"):
        return redirect(url_for("superadmin_login"))
        
    if request.method == "POST":
        # Handle Tenant Creation
        name = request.form.get("name")
        subdomain = request.form.get("subdomain")
        color = request.form.get("color", "#f16924")
        
        if not name or not subdomain:
            flash("Name and Subdomain are required", "error")
        else:
            success = provision_new_tenant(name, subdomain, color)
            if success:
                flash(f"Tenant '{name}' created successfully! URL: http://{subdomain}.plotpro.in", "success")
            else:
                flash("Failed to create tenant. Check logs.", "error")
    
    # Fetch all tenants
    master_conn = mysql.connector.connect(
        host=os.getenv("DB_HOST", "localhost"),
        user=os.getenv("DB_USER", "root"),
        password=os.getenv("DB_PASSWORD", ""),
        database="plotpro_master"
    )
    c = master_conn.cursor(dictionary=True)
    c.execute("SELECT * FROM tenants ORDER BY created_at DESC")
    tenants = c.fetchall()
    
    # Fetch all leads
    c.execute("SELECT * FROM leads ORDER BY created_at DESC")
    leads = c.fetchall()
    
    master_conn.close()
    
    return render_template("superadmin_dashboard.html", tenants=tenants, leads=leads)

@app.route("/superadmin/delete_tenant/<int:tenant_id>", methods=["POST"])
def delete_tenant(tenant_id):
    if not session.get("is_super_admin"):
        return redirect(url_for("superadmin_login"))
    
    try:
        master_conn = mysql.connector.connect(
            host=os.getenv("DB_HOST", "localhost"),
            user=os.getenv("DB_USER", "root"),
            password=os.getenv("DB_PASSWORD", ""),
            database="plotpro_master"
        )
        c = master_conn.cursor(dictionary=True)
        
        # Get DB Name first
        c.execute("SELECT db_name, name FROM tenants WHERE id=%s", (tenant_id,))
        tenant = c.fetchone()
        
        if tenant:
            # 1. Drop the Tenant Database (Crucial for cleanup)
            # CAREFUL: This is destructive!
            try:
                c.execute(f"DROP DATABASE IF EXISTS `{tenant['db_name']}`")
            except Exception as db_err:
                print(f"Error dropping DB {tenant['db_name']}: {db_err}")
                # Proceed to remove from master anyway
            
            # 2. Remove from Master
            c.execute("DELETE FROM tenants WHERE id=%s", (tenant_id,))
            master_conn.commit()
            flash(f"Tenant '{tenant['name']}' and database deleted.", "success")
        else:
            flash("Tenant not found.", "error")
            
        master_conn.close()
    except Exception as e:
        flash(f"Error deleting tenant: {e}", "error")
        
    return redirect(url_for("superadmin_dashboard"))

@app.route("/superadmin/update_tenant/<int:tenant_id>", methods=["POST"])
def update_tenant(tenant_id):
    if not session.get("is_super_admin"):
        return redirect(url_for("superadmin_login"))
        
    name = request.form.get("name")
    subdomain = request.form.get("subdomain")
    color = request.form.get("color")
    
    try:
        master_conn = mysql.connector.connect(
            host=os.getenv("DB_HOST", "localhost"),
            user=os.getenv("DB_USER", "root"),
            password=os.getenv("DB_PASSWORD", ""),
            database="plotpro_master"
        )
        c = master_conn.cursor()
        
        # We allow changing subdomain. This maps New Subdomain -> Old DB Name.
        # This effectively renames the site URL without moving data.
        c.execute("""
            UPDATE tenants 
            SET name=%s, subdomain=%s, brand_color=%s 
            WHERE id=%s
        """, (name, subdomain, color, tenant_id))
        
        master_conn.commit()
        master_conn.close()
        flash(f"Tenant '{name}' updated successfully.", "success")
    except Exception as e:
        flash(f"Error updating tenant: {e}", "error")
        
    return redirect(url_for("superadmin_dashboard"))


@app.route("/create", methods=["POST"])
def create():
    form = request.form
    no = form.get("no", "").strip()
    project_name = form.get("project_name", "").strip()
    date = form.get("date", "").strip()
    venture = form.get("venture", "").strip()
    customer_name = form.get("customer_name", "").strip()

    amount_numeric = form.get("amount_numeric", "0").replace(",", "").strip()
    amount_words = form.get("amount_words", "").strip() or number_to_words(amount_numeric)

    plot_no = form.get("plot_no", "").strip()
    square_yards = form.get("square_yards", "").strip()
    purpose = form.get("purpose", "").strip()
    drawn_bank = form.get("drawn_bank", "").strip()
    branch = form.get("branch", "").strip()
    payment_mode = form.get("payment_mode", "").strip()
    instrument_no = form.get("instrument_no", "").strip()
    
    # New fields
    pan_no = form.get("pan_no", "").strip().upper()
    aadhar_no = form.get("aadhar_no", "").strip()
    basic_price = form.get("basic_price", "").replace(",", "").strip()

    created_at = datetime.utcnow().isoformat()
    
    # Check if user is admin
    is_admin = session.get("role") == "admin"

    conn = database.get_db_connection()
    c = conn.cursor()
    
    if is_admin:
        # Admin: Save directly to receipts table
        c.execute(
            """
            INSERT INTO receipts
            (no, project_name, date, venture, customer_name, amount_numeric, amount_words,
             plot_no, square_yards, purpose, drawn_bank, branch, payment_mode, instrument_no, 
             pan_no, aadhar_no, basic_price, created_at)
            VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
        """,
            (
                no,
                project_name,
                date,
                venture,
                customer_name,
                amount_numeric,
                amount_words,
                plot_no,
                square_yards,
                purpose,
                drawn_bank,
                branch,
                payment_mode,
                instrument_no,
                pan_no,
                aadhar_no,
                basic_price,
                created_at,
            ),
        )
        rid = c.lastrowid
        
        # If basic_price is provided, update it for all other receipts of this plot
        if basic_price and plot_no and project_name:
            try:
                c.execute(
                    "UPDATE receipts SET basic_price = %s WHERE plot_no = %s AND project_name = %s",
                    (basic_price, plot_no, project_name),
                )
            except database.OperationalError:
                pass
                
        conn.commit()
        conn.close()
        flash("Receipt created successfully!", "success")
        return redirect(url_for("view_receipt", receipt_id=rid))
    else:
        # Non-admin: Save to pending_receipts for approval
        # Note: pending_receipts table might not have new columns yet. 
        # We save what we can. If strict schema, this needs migration. 
        # For now, sticking to existing columns for pending to avoid errors if migration not run.
        # User is admin in this context so this block is less critical but should be updated eventually.
        c.execute(
            """
            INSERT INTO pending_receipts
            (no, project_name, date, venture, customer_name, amount_numeric, amount_words,
             plot_no, square_yards, purpose, drawn_bank, branch, payment_mode, instrument_no,
             submitted_by, submitted_at, status)
            VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, 'pending')
        """,
            (
                no,
                project_name,
                date,
                venture,
                customer_name,
                amount_numeric,
                amount_words,
                plot_no,
                square_yards,
                purpose,
                drawn_bank,
                branch,
                payment_mode,
                instrument_no,
                session.get("username", "unknown"),
                created_at,
            ),
        )
        conn.commit()
        conn.close()
        flash("Receipt submitted for admin approval. You will be notified once approved.", "info")
        return redirect(url_for("index"))


@app.route("/receipt/<int:receipt_id>")
def view_receipt(receipt_id):
    conn = database.get_db_connection()
    c = conn.cursor()
    c.execute("SELECT * FROM receipts WHERE id = %s", (receipt_id,))
    row = database.fetch_one(c)
    conn.close()
    if not row:
        abort(404)
    r = dict_from_row(row)
    r["amount_formatted"] = format_inr(r["amount_numeric"])
    return render_template("receipt_boot.html", r=r, pdf_mode=False)


@app.route("/receipt/<int:receipt_id>/view_html")
def receipt_view_html(receipt_id):
    conn = database.get_db_connection()
    c = conn.cursor()
    c.execute("SELECT * FROM receipts WHERE id = %s", (receipt_id,))
    row = database.fetch_one(c)
    conn.close()
    if not row:
        abort(404)
    r = dict_from_row(row)
    r["amount_formatted"] = format_inr(r["amount_numeric"])
    return render_template("receipt_boot.html", r=r, pdf_mode=False)


@app.route("/receipt/<int:receipt_id>/edit")
def edit_receipt(receipt_id):
    """
    Render the form.html in edit mode for the given receipt ID.
    """
    conn = database.get_db_connection()
    c = conn.cursor()
    c.execute("SELECT * FROM receipts WHERE id = %s", (receipt_id,))
    row = database.fetch_one(c)
    conn.close()
    if not row:
        abort(404)
    r = dict_from_row(row)
    idx = get_column_index("receipts", "basic_price")
    if idx is not None and idx < len(row):
        r["basic_price"] = row[idx]
    projects = get_projects()
    return render_template("form.html", recent=[], r=r, edit_mode=True, projects=projects)


@app.route("/receipt/<int:receipt_id>/update", methods=["POST"])
def update_receipt(receipt_id):
    form = request.form
    no = form.get("no", "").strip()
    project_name = form.get("project_name", "").strip()
    date = form.get("date", "").strip()
    venture = form.get("venture", "").strip()
    customer_name = form.get("customer_name", "").strip()

    amount_numeric = form.get("amount_numeric", "0").replace(",", "").strip()
    amount_words = form.get("amount_words", "").strip() or number_to_words(amount_numeric)

    plot_no = form.get("plot_no", "").strip()
    square_yards = form.get("square_yards", "").strip()
    purpose = form.get("purpose", "").strip()
    drawn_bank = form.get("drawn_bank", "").strip()
    branch = form.get("branch", "").strip()
    payment_mode = form.get("payment_mode", "").strip()
    instrument_no = form.get("instrument_no", "").strip()
    basic_price = form.get("basic_price", "").replace(",", "").strip()

    conn = database.get_db_connection()
    c = conn.cursor()
    c.execute(
        """
        UPDATE receipts SET
            no=%s, project_name=%s, date=%s, venture=%s, customer_name=%s,
            amount_numeric=%s, amount_words=%s, plot_no=%s, square_yards=%s,
            purpose=%s, drawn_bank=%s, branch=%s, payment_mode=%s, instrument_no=%s, basic_price=%s
        WHERE id=%s
    """,
        (
            no,
            project_name,
            date,
            venture,
            customer_name,
            amount_numeric,
            amount_words,
            plot_no,
            square_yards,
            purpose,
            drawn_bank,
            branch,
            payment_mode,
            instrument_no,
            basic_price,
            receipt_id,
        ),
    )
    if basic_price and plot_no:
        try:
            c.execute(
                "UPDATE receipts SET basic_price = %s WHERE plot_no = %s",
                (basic_price, plot_no),
            )
        except database.OperationalError:
            pass
    conn.commit()
    conn.close()
    return redirect(url_for("view_receipt", receipt_id=receipt_id))


@app.route("/receipt/<int:receipt_id>/delete", methods=["POST"])
def delete_receipt(receipt_id):
    if session.get("role") != "admin":
        return jsonify({"success": False, "error": "Unauthorized"}), 403
        
    conn = database.get_db_connection()
    c = conn.cursor()
    try:
        c.execute("DELETE FROM receipts WHERE id = %s", (receipt_id,))
        conn.commit()
        return jsonify({"success": True})
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500
    finally:
        conn.close()


# -------------------------
# Playwright-based PDF generation
# -------------------------
def html_to_pdf_bytes_playwright(html_string, base_url=None, landscape=False):
    """
    Use Playwright to render HTML to PDF bytes.
    Injects a relaxed CSP meta tag and a <base> element (for static paths) only inside Playwright
    so the production site CSP is not modified.
    """
    if sync_playwright is None:
        raise RuntimeError(
            "Playwright not available. Install using:\n"
            "pip install playwright\n"
            "python -m playwright install chromium"
        )

    with sync_playwright() as pw:
        browser = pw.chromium.launch()
        page = browser.new_page()

        # Inject a meta CSP that allows eval/inline for the headless context only
        # (keeps your public site CSP unchanged)
        csp_script = """
        (() => {
          try {
            var meta = document.createElement('meta');
            meta.httpEquiv = "Content-Security-Policy";
            meta.content = "default-src * 'unsafe-inline' 'unsafe-eval' data: blob:; script-src * 'unsafe-inline' 'unsafe-eval' data: blob:; style-src * 'unsafe-inline' data:; img-src * data: blob:;";
            var h = document.getElementsByTagName('head')[0] || document.documentElement;
            h.appendChild(meta);
          } catch (e) {
            // ignore
          }
        })();
        """
        page.add_init_script(csp_script)

        # set page content
        page.set_content(html_string, wait_until="networkidle")

        # If base_url provided, add a <base> tag (helps resolving relative static links)
        if base_url:
            try:
                page.evaluate(
                    """(b) => {
                          try {
                              var base = document.createElement('base');
                              base.href = b;
                              var h = document.getElementsByTagName('head')[0] || document.documentElement;
                              // insert front so relative URLs resolve
                              if (h.firstChild) h.insertBefore(base, h.firstChild); else h.appendChild(base);
                          } catch (e) {}
                      }""",
                    base_url,
                )
            except Exception:
                # non-fatal
                pass

        pdf_bytes = page.pdf(format="A4", print_background=True)
        browser.close()
        return pdf_bytes


def render_receipt_pdf_bytes(r):
    logo_path = Path(app.static_folder) / "images" / "logo.png"
    logo_data = None

    if logo_path.exists():
        with open(logo_path, "rb") as f:
            b = f.read()
            mime = (
                "image/jpeg" if logo_path.suffix.lower() in [".jpg", ".jpeg"]
                else "image/svg+xml" if logo_path.suffix.lower() == ".svg"
                else "image/png"
            )
            logo_data = f"data:{mime};base64,{base64.b64encode(b).decode('utf8')}"

    r = dict(r)
    r["logo_data"] = logo_data

    html_string = render_template("receipt_boot.html", r=r, pdf_mode=True)
    base_url = os.path.abspath(".")
    # call Playwright wrapper
    return html_to_pdf_bytes_playwright(html_string, base_url=base_url)


@app.route("/receipt/<int:receipt_id>/pdf")
def receipt_pdf(receipt_id):
    # fetch receipt
    conn = database.get_db_connection()
    c = conn.cursor()
    c.execute("SELECT * FROM receipts WHERE id = %s", (receipt_id,))
    row = database.fetch_one(c)
    conn.close()
    if not row:
        abort(404)

    r = dict_from_row(row)
    r["amount_formatted"] = format_inr(r["amount_numeric"])

    # render PDF bytes
    try:
        pdf_bytes = render_receipt_pdf_bytes(r)
    except Exception as e:
        app.logger.exception("Error rendering PDF for receipt %s", receipt_id)
        abort(500, description=f"Error generating PDF: {e}")

    # ensure bytes
    if isinstance(pdf_bytes, memoryview):
        pdf_bytes = bytes(pdf_bytes)
    if not isinstance(pdf_bytes, (bytes, bytearray)):
        app.logger.error("render_receipt_pdf_bytes returned non-bytes for receipt %s: %r", receipt_id, type(pdf_bytes))
        abort(500, description="PDF generator returned unexpected type")

    # sanity check (optional but helpful)
    if not pdf_bytes.startswith(b"%PDF"):
        app.logger.error("Generated content does not look like a PDF: %r", pdf_bytes[:32])
        abort(500, description="PDF generation failed (invalid output)")

    # Build a friendly filename using receipt 'no' if present, else fallback to id.
    raw_name = (r.get("no") or str(receipt_id)).strip()
    # sanitize: produce safe fragment
    frag = _safe_filename_fragment(raw_name, str(receipt_id))
    filename = f"receipt_{frag}.pdf"

    buf = io.BytesIO(pdf_bytes)
    buf.seek(0)

    # Prefer send_file with download_name (Flask >= 2.2). Fallback to older param.
    try:
        response = send_file(
            buf,
            mimetype="application/pdf",
            as_attachment=True,
            download_name=filename,
            conditional=False
        )
    except TypeError:
        response = send_file(
            buf,
            mimetype="application/pdf",
            as_attachment=True,
            attachment_filename=filename,
            conditional=False
        )

    # Add a robust Content-Disposition header (filename + filename* UTF-8)
    try:
        # urlquote filename for filename* parameter
        quoted = urlquote(filename)
        response.headers["Content-Disposition"] = f'attachment; filename="{filename}"; filename*=UTF-8\'\'{quoted}'
    except Exception:
        # ignore if we cannot set
        pass

    # Explicit helpful headers
    response.headers["Cache-Control"] = "no-store, no-cache, must-revalidate, post-check=0, pre-check=0"
    response.headers["Pragma"] = "no-cache"
    try:
        response.headers["Content-Length"] = str(len(pdf_bytes))
    except Exception:
        pass

    return response


# -----------------------------
# Search receipts by plot (HTML)
# -----------------------------
@app.route("/search_by_plot", methods=["GET", "POST"])
def search_by_plot():
    # Check if user has permission to search receipts
    if not (session.get("role") == "admin" or session.get("can_search_receipts")):
        abort(403)
    
    plot_no = None
    project_name = None
    results = None
    projects = get_projects()

    if request.method == "POST":
        plot_no = (request.form.get("plot_no") or "").strip()
        project_name = (request.form.get("project_name") or "").strip()
        
        if plot_no:
            conn = database.get_db_connection()
            c = conn.cursor()

            query = "SELECT * FROM receipts WHERE LOWER(TRIM(plot_no)) = LOWER(TRIM(%s))"
            params = [plot_no]
            
            if project_name:
                query += " AND project_name = %s"
                params.append(project_name)
                
            query += " ORDER BY id DESC"
            
            c.execute(query, tuple(params))
            rows = database.fetch_all(c)

            if not rows:
                like = f"%{plot_no}%"
                query = "SELECT * FROM receipts WHERE plot_no LIKE %s"
                params = [like]
                if project_name:
                    query += " AND project_name = %s"
                    params.append(project_name)
                query += " ORDER BY id DESC"
                
                c.execute(query, tuple(params))
                rows = database.fetch_all(c)

            conn.close()

            results = []
            for row in rows:
                r = dict_from_row(row)
                r["amount_formatted"] = format_inr(r["amount_numeric"])
                results.append(r)

    return render_template("search_by_plot.html", results=results, plot_no=plot_no, projects=projects, selected_project=project_name)


# -----------------------------
# API Endpoint (JSON lookup)
# -----------------------------
@app.route("/api/plot_lookup")
def plot_lookup():
    plot_no = request.args.get("plot_no", "").strip()
    project_name = request.args.get("project_name", "").strip()
    
    if not plot_no:
        return jsonify({"found": False, "error": "plot_no required"}), 400
    
    if not project_name:
        return jsonify({"found": False, "error": "project_name required"}), 400

    conn = database.get_db_connection()
    c = conn.cursor()
    
    # Lookup using both plot_no AND project_name as unique key
    c.execute("SELECT * FROM receipts WHERE plot_no = %s AND project_name = %s ORDER BY id DESC LIMIT 1", (plot_no, project_name))
        
    row = database.fetch_one(c)
    conn.close()

    if not row:
        return jsonify({"found": False}), 200

    r = dict_from_row(row)
    r["amount_formatted"] = format_inr(r["amount_numeric"])
    idx = get_column_index("receipts", "basic_price")
    if idx is not None and idx < len(row):
        r["basic_price"] = row[idx]
    return jsonify({"found": True, **r})


# -----------------------------
# Dashboard & Analytics
# -----------------------------
@app.route("/dashboard")
def dashboard():
    """Display project dashboard with statistics"""
    # If user doesn't have dashboard access but has Vishvam access, redirect to Vishvam Layout
    if not (session.get("role") == "admin" or session.get("can_view_dashboard")):
        if session.get("can_view_vishvam_layout"):
            return redirect(url_for("plot_layout_viewer", project_name="Vishvam"))
        abort(403)
    selected_project = request.args.get("project", "").strip()
    
    conn = database.get_db_connection()
    c = conn.cursor()
    
    # Get all projects
    projects = get_projects()
    
    # Build query based on filter
    if selected_project:
        # Count unique plots sold for this project
        c.execute("""
            SELECT COUNT(DISTINCT plot_no) 
            FROM receipts 
            WHERE project_name = %s AND plot_no IS NOT NULL AND plot_no != ''
        """, (selected_project,))
        plots_sold_row = database.fetch_one(c)
        num_plots = plots_sold_row[0] if plots_sold_row and plots_sold_row[0] is not None else 0
        
        # Get total plots and landowner plots for this project
        c.execute("SELECT total_plots, plots_to_landowners FROM projects WHERE name = %s", (selected_project,))
        row = database.fetch_one(c)
        total_plots = row[0] if row else 0
        plots_to_landowners = row[1] if row and len(row) > 1 else 0
    else:
        # Count all unique plots sold
        c.execute("""
            SELECT COUNT(DISTINCT plot_no) 
            FROM receipts 
            WHERE plot_no IS NOT NULL AND plot_no != ''
        """)
        num_plots = database.fetch_one(c)[0] or 0
        
        # Sum total plots and landowner plots from all projects
        c.execute("SELECT SUM(total_plots), SUM(plots_to_landowners) FROM projects")
        row = database.fetch_one(c)
        total_plots = row[0] if row and row[0] else 0
        plots_to_landowners = row[1] if row and row[1] else 0
    
    conn.close()
    
    # Get pending receipts count for admin notification
    pending_count = 0
    if session.get("role") == "admin":
        conn2 = database.get_db_connection()
        c2 = conn2.cursor()
        c2.execute("SELECT COUNT(*) FROM pending_receipts WHERE status = 'pending'")
        row = database.fetch_one(c2)
        pending_count = row[0] if row else 0
        conn2.close()
    
    plots_remaining = max(0, total_plots - plots_to_landowners - num_plots)
    
    return render_template(
        "dashboard.html",
        projects=projects,
        selected_project=selected_project,
        total_plots=total_plots,
        num_plots=num_plots,
        plots_remaining=plots_remaining,
        plots_to_landowners=plots_to_landowners,
        pending_count=pending_count
    )


@app.route("/api/stats/amount_by_month")
def stats_amount_by_month():
    """Return total receipt amount grouped by month (YYYY-MM) for Google Charts.

    Respects optional ?project=<name> filter to stay in sync with dashboard selection.
    """
    if not (session.get("role") == "admin" or session.get("can_view_dashboard")):
        abort(403)

    selected_project = request.args.get("project", "").strip()

    conn = database.get_db_connection()
    c = conn.cursor()

    query = """
        SELECT strftime('%Y-%m', date) AS ym, COALESCE(SUM(amount_numeric), 0.0)
        FROM receipts
        WHERE date IS NOT NULL AND TRIM(date) != ''
    """
    params = []

    if selected_project:
        query += " AND project_name = %s"
        params.append(selected_project)

    query += " GROUP BY ym ORDER BY ym ASC"

    c.execute(query, params)
    rows = database.fetch_all(c)
    conn.close()

    data = [
        {
            "month": row[0] or "Unknown",
            "total": float(row[1] or 0.0),
        }
        for row in rows
    ]

    return jsonify(data)


@app.route("/api/stats/amount_by_project")
def stats_amount_by_project():
    """Return total receipt amount grouped by project for Google Charts.

    If a specific project is selected on the dashboard, this endpoint still returns
    only that project's total (useful for future per-project breakdowns), but
    when no project is selected it returns all projects.
    """
    if not (session.get("role") == "admin" or session.get("can_view_dashboard")):
        abort(403)

    selected_project = request.args.get("project", "").strip()

    conn = database.get_db_connection()
    c = conn.cursor()

    query = """
        SELECT
            COALESCE(NULLIF(TRIM(project_name), ''), 'Unknown') AS project_name,
            COALESCE(SUM(amount_numeric), 0.0) AS total
        FROM receipts
        WHERE 1 = 1
    """
    params = []

    if selected_project:
        query += " AND project_name = %s"
        params.append(selected_project)

    query += " GROUP BY project_name ORDER BY total DESC"

    c.execute(query, params)
    rows = database.fetch_all(c)
    conn.close()

    data = [
        {
            "project": row[0] or "Unknown",
            "total": float(row[1] or 0.0),
        }
        for row in rows
    ]

    return jsonify(data)


@app.route("/api/stats/amount_by_payment_mode")
def stats_amount_by_payment_mode():
    """Return total receipt amount grouped by payment_mode for Google Charts.

    Respects optional ?project=<name> filter.
    """
    if not (session.get("role") == "admin" or session.get("can_view_dashboard")):
        abort(403)

    selected_project = request.args.get("project", "").strip()

    conn = database.get_db_connection()
    c = conn.cursor()

    query = """
        SELECT
            COALESCE(NULLIF(TRIM(payment_mode), ''), 'Unknown') AS mode,
            COALESCE(SUM(amount_numeric), 0.0) AS total
        FROM receipts
        WHERE 1 = 1
    """
    params = []

    if selected_project:
        query += " AND project_name = %s"
        params.append(selected_project)

    query += " GROUP BY mode ORDER BY total DESC"

    c.execute(query, params)
    rows = database.fetch_all(c)
    conn.close()

    data = [
        {
            "mode": row[0] or "Unknown",
            "total": float(row[1] or 0.0),
        }
        for row in rows
    ]

    return jsonify(data)


@app.route("/api/stats/commission_by_project")
def stats_commission_by_project():
    """Return total commission grouped by project for Google Charts.
    
    Respects optional %sproject= filter.
    """
    if not (session.get("role") == "admin" or session.get("can_view_dashboard")):
        abort(403)

    selected_project = request.args.get("project", "").strip()

    conn = database.get_db_connection()
    c = conn.cursor()

    query = """
        SELECT
            COALESCE(NULLIF(TRIM(project_name), ''), 'Unknown') AS project_name,
            COALESCE(SUM(cgm_rate + srgm_rate + gm_rate), 0.0) AS total
        FROM commissions
        WHERE 1 = 1
    """
    params = []

    if selected_project:
        query += " AND project_name = %s"
        params.append(selected_project)

    query += " GROUP BY project_name ORDER BY total DESC"

    c.execute(query, params)
    rows = database.fetch_all(c)
    conn.close()

    data = [
        {
            "project": row[0] or "Unknown",
            "total": float(row[1] or 0.0),
        }
        for row in rows
    ]

    return jsonify(data)





@app.route("/api/stats/projects_inventory")
def stats_projects_inventory():
    """Return projects with total plots and sold plots for Google Charts.
    
    If %sproject= filter is set, returns just that project.
    """
    if not (session.get("role") == "admin" or session.get("can_view_dashboard")):
        abort(403)

    selected_project = request.args.get("project", "").strip()

    conn = database.get_db_connection()
    c = conn.cursor()

    data = []
    if selected_project:
        # Get total plots for this project
        c.execute("SELECT total_plots FROM projects WHERE name = %s", (selected_project,))
        proj_row = database.fetch_one(c)
        total_plots = proj_row[0] if proj_row else 0

        # Count sold plots for this project
        c.execute("""
            SELECT COUNT(DISTINCT plot_no)
            FROM receipts
            WHERE project_name = %s AND plot_no IS NOT NULL AND plot_no != ''
        """, (selected_project,))
        sold_row = database.fetch_one(c)
        sold_plots = sold_row[0] if sold_row else 0

        data.append({
            "project": selected_project,
            "total_plots": int(total_plots or 0),
            "sold_plots": int(sold_plots or 0),
            "available_plots": max(0, int(total_plots or 0) - int(sold_plots or 0))
        })
    else:
        # All projects
        c.execute("SELECT name, total_plots FROM projects ORDER BY name")
        proj_rows = database.fetch_all(c)
        for proj_name, total_plots in proj_rows:
            # Count sold plots for this project
            c.execute("""
                SELECT COUNT(DISTINCT plot_no)
                FROM receipts
                WHERE project_name = %s AND plot_no IS NOT NULL AND plot_no != ''
            """, (proj_name,))
            sold_row = database.fetch_one(c)
            sold_plots = sold_row[0] if sold_row else 0

            data.append({
                "project": proj_name or "Unknown",
                "total_plots": int(total_plots or 0),
                "sold_plots": int(sold_plots or 0),
                "available_plots": max(0, int(total_plots or 0) - int(sold_plots or 0))
            })

    conn.close()
    return jsonify(data)


@app.route("/api/stats/cgm_plot_sales")
def stats_cgm_plot_sales():
    """Return number of plots sold by each CGM for Google Charts.
    
    Joins receipts with commissions to get CGM per receipt.
    Respects optional %sproject= filter.
    """
    if not (session.get("role") == "admin" or session.get("can_view_dashboard")):
        abort(403)

    selected_project = request.args.get("project", "").strip()

    conn = database.get_db_connection()
    c = conn.cursor()

    query = """
        SELECT
            COALESCE(c.cgm_name, 'Unknown') AS cgm_name,
            COUNT(DISTINCT r.plot_no) AS plots_sold
        FROM receipts r
        LEFT JOIN commissions c ON r.project_name = c.project_name AND r.plot_no = c.plot_no
        WHERE r.plot_no IS NOT NULL AND r.plot_no != ''
    """
    params = []

    if selected_project:
        query += " AND r.project_name = %s"
        params.append(selected_project)

    query += " GROUP BY cgm_name ORDER BY plots_sold DESC"

    c.execute(query, params)
    rows = database.fetch_all(c)
    conn.close()

    data = [
        {
            "cgm": row[0] or "Unknown",
            "plots_sold": int(row[1] or 0),
        }
        for row in rows
    ]

    return jsonify(data)





@app.route("/debug/tables")
def debug_tables():
    """Print all tables in receipts.db to console for debugging."""
    conn = database.get_db_connection()
    c = conn.cursor()
    c.execute("SELECT name FROM sqlite_master WHERE type='table' ORDER BY name")
    tables = database.fetch_all(c)
    conn.close()

    print("=== Tables in receipts.db ===")
    for table in tables:
        print(table[0])
    print("==============================")

    return "<pre>" + "\n".join(t[0] for t in tables) + "</pre>"


@app.route("/analytics")
def analytics():
    if not (session.get("role") == "admin" or session.get("can_view_dashboard")):
        abort(403)

    selected_project = request.args.get("project", "").strip()

    conn = database.get_db_connection()
    c = conn.cursor()

    projects = get_projects()

    # Pending receipts count for admin
    pending_count = 0
    if session.get("role") == "admin":
        c.execute("SELECT COUNT(*) FROM pending_receipts WHERE status = 'pending'")
        row = database.fetch_one(c)
        pending_count = row[0] if row else 0

    conn.close()

    return render_template(
        "analytics.html",
        projects=projects,
        selected_project=selected_project,
        pending_count=pending_count,
    )


@app.route("/plot_management", methods=['GET', 'POST'])
def plot_management():
    """Manage plot configurations for projects"""
    if not session.get("role") == "admin":
        if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
            return jsonify({'error': 'Access denied. Admin privileges required.'}), 403
        flash("Access denied. Admin privileges required.", "danger")
        return redirect(url_for('dashboard'))
    
    projects = get_projects()
    project_data = None
    selected_project = request.form.get('project') or request.args.get('project')
    
    if request.method == 'POST':
        selected_project = request.form.get('project')
        total_plots = request.form.get('total_plots', type=int)
        plots_to_landowners = request.form.get('plots_to_landowners', type=int)
        
        if selected_project and total_plots is not None and plots_to_landowners is not None:
            conn = database.get_db_connection()
            c = conn.cursor()
            
            # Check if project exists
            c.execute("SELECT id FROM projects WHERE name = %s", (selected_project,))
            project = database.fetch_one(c)
            
            # Get number of plots sold for this project
            c.execute("""
                SELECT COUNT(DISTINCT plot_no) 
                FROM receipts 
                WHERE project_name = %s AND plot_no IS NOT NULL AND plot_no != ''
            """, (selected_project,))
            plots_sold_row = database.fetch_one(c)
            plots_sold = plots_sold_row[0] if plots_sold_row and plots_sold_row[0] is not None else 0
            
            # Calculate available plots
            plots_available = max(0, total_plots - plots_to_landowners - plots_sold)
            
            if project:
                # Update existing project
                c.execute("""
                    UPDATE projects 
                    SET total_plots = %s, plots_to_landowners = %s
                    WHERE name = %s
                """, (total_plots, plots_to_landowners, selected_project))
                flash(f"Updated {selected_project}: {total_plots} total plots, {plots_to_landowners} to landowners", "success")
            else:
                # Insert new project
                c.execute("""
                    INSERT INTO projects (name, total_plots, plots_to_landowners)
                    VALUES (%s, %s, %s)
                """, (selected_project, total_plots, plots_to_landowners))
                flash(f"Added {selected_project}: {total_plots} total plots, {plots_to_landowners} to landowners", "success")
            
            conn.commit()
            conn.close()
        else:
            flash("Please fill in all fields", "warning")
    
    # If project is selected (either via POST or GET), load its data
    if selected_project:
        conn = database.get_db_connection()
        c = conn.cursor()
        
        # Get project details
        c.execute("SELECT total_plots, plots_to_landowners FROM projects WHERE name = %s", (selected_project,))
        row = database.fetch_one(c)
        
        # Get number of plots sold for this project
        c.execute("""
            SELECT COUNT(DISTINCT plot_no) 
            FROM receipts 
            WHERE project_name = %s AND plot_no IS NOT NULL AND plot_no != ''
        """, (selected_project,))
        plots_sold_row = database.fetch_one(c)
        plots_sold = plots_sold_row[0] if plots_sold_row and plots_sold_row[0] is not None else 0
        
        if row:
            total_plots = row['total_plots'] or 0
            plots_to_landowners = row['plots_to_landowners'] or 0
            plots_available = max(0, total_plots - plots_to_landowners - plots_sold)
            
            project_data = {
                'total_plots': total_plots,
                'plots_to_landowners': plots_to_landowners,
                'plots_available': plots_available,
                'plots_sold': plots_sold
            }
        conn.close()
    
    # Return JSON if requested via AJAX
    if request.headers.get('X-Requested-With') == 'XMLHttpRequest' or request.args.get('json'):
        return jsonify({
            'project_data': project_data,
            'projects': projects,
            'selected_project': selected_project
        })
    

    
    return render_template(
        "plot_management.html",
        project_data=project_data,
        selected_project=selected_project,
        projects=projects
    )

@app.route("/update_plot_settings", methods=["POST"])
def update_plot_settings():
    if not session.get("role") == "admin":
        return redirect(url_for('login'))
        
    project_name = request.form.get("project_name")
    try:
        total_plots = int(request.form.get("total_plots", 0))
        plots_to_landowners = int(request.form.get("plots_to_landowners", 0))
    except ValueError:
        flash("Invalid input values", "danger")
        return redirect(url_for('plot_management'))
        
    conn = database.get_db_connection()
    c = conn.cursor()
    try:
        # Check if project exists in projects table, if not create/ignore?
        # Assuming project exists from context
        c.execute("UPDATE projects SET total_plots = %s, plots_to_landowners = %s WHERE name = %s", 
                 (total_plots, plots_to_landowners, project_name))
        conn.commit()
        flash("Plot settings updated successfully", "success")
    except Exception as e:
        flash(f"Error updating settings: {str(e)}", "danger")
    finally:
        conn.close()
        
    return redirect(url_for('plot_management'))

@app.route("/layout_management")
def layout_management():
    if not session.get("role") == "admin":
        return redirect(url_for('login'))
        
    # Get all projects with full details for layout_management.html
    conn = database.get_db_connection()
    c = conn.cursor()
    c.execute("SELECT * FROM projects ORDER BY name")
    all_projects_rows = database.fetch_all(c)
    conn.close()
    
    active_projects = []
    archived_projects = []
    
    for p in all_projects_rows:
        p_dict = dict(p)
        if p_dict.get('is_archived'):
            archived_projects.append(p_dict)
        else:
            active_projects.append(p_dict)
            
    return render_template(
        "layout_management.html",
        active_projects=active_projects,
        archived_projects=archived_projects
    )


# -----------------------------
# Project Management Routes
# -----------------------------
@app.route("/projects/add", methods=["POST"])
def add_project():
    """Add a new project"""
    if not session.get("role") == "admin":
        flash("Access denied. Admin privileges required.", "danger")
        return redirect(url_for('plot_management'))
        
    project_name = request.form.get("project_name", "").strip()
    
    if not project_name:
        flash("Project name is required", "warning")
        return redirect(url_for('plot_management'))
        
    try:
        conn = database.get_db_connection()
        c = conn.cursor()
        
        # Check if project already exists
        c.execute("SELECT id FROM projects WHERE name = %s", (project_name,))
        if c.fetchone():
            flash(f"Project '{project_name}' already exists", "warning")
            conn.close()
            return redirect(url_for('plot_management'))
            
        # Insert new project
        c.execute("INSERT INTO projects (name, total_plots, plots_to_landowners) VALUES (%s, 0, 0)", (project_name,))
        conn.commit()
        conn.close()
        
        flash(f"Project '{project_name}' created successfully", "success")
        
    except Exception as e:
        app.logger.error(f"Error adding project: {e}")
        flash("An error occurred while creating the project", "danger")
        
    return redirect(url_for('plot_management'))


@app.route("/projects/update", methods=["POST"])
def update_project():
    """Update project name"""
    if not session.get("role") == "admin":
        flash("Access denied. Admin privileges required.", "danger")
        return redirect(url_for('plot_management'))
        
    project_id = request.form.get("project_id")
    new_name = request.form.get("new_name", "").strip()
    
    if not project_id or not new_name:
        flash("Project ID and new name are required", "warning")
        return redirect(url_for('plot_management'))
        
    try:
        conn = database.get_db_connection()
        c = conn.cursor()
        
        # Check if new name already taken by OTHER project
        c.execute("SELECT id FROM projects WHERE name = %s AND id != %s", (new_name, project_id))
        if c.fetchone():
            flash(f"Project name '{new_name}' is already in use", "warning")
            conn.close()
            return redirect(url_for('plot_management'))
            
        # Update name
        c.execute("UPDATE projects SET name = %s WHERE id = %s", (new_name, project_id))
        
        # Also update receipts table? (Since receipts store project_name...)
        # Ideally we should use foreign keys, but currently it seems to store name string.
        # So we probably need to update receipts too to keep data consistent.
        c.execute("UPDATE receipts SET project_name = %s WHERE project_name = (SELECT name FROM projects WHERE id = %s)", 
                 (new_name, project_id))
                 
        conn.commit()
        conn.close()
        
        flash(f"Project updated to '{new_name}' successfully", "success")
        
    except Exception as e:
        app.logger.error(f"Error updating project: {e}")
        flash("An error occurred while updating the project", "danger")
        
    return redirect(url_for('plot_management'))


@app.route("/projects/archive", methods=["POST"])
def archive_project():
    """Archive a project"""
    if not session.get("role") == "admin":
        flash("Access denied. Admin privileges required.", "danger")
        return redirect(url_for('plot_management'))
        
    project_id = request.form.get("project_id")
    
    if not project_id:
        flash("Project ID is required", "warning")
        return redirect(url_for('plot_management'))
        
    try:
        conn = database.get_db_connection()
        c = conn.cursor()
        
        c.execute("UPDATE projects SET is_archived = 1 WHERE id = %s", (project_id,))
        conn.commit()
        conn.close()
        
        flash("Project archived successfully", "success")
        
    except Exception as e:
        app.logger.error(f"Error archiving project: {e}")
        flash("An error occurred while archiving the project", "danger")
        
    return redirect(url_for('plot_management'))


@app.route("/projects/unarchive", methods=["POST"])
def unarchive_project():
    """Restores an archived project"""
    if not session.get("role") == "admin":
        flash("Access denied. Admin privileges required.", "danger")
        return redirect(url_for('layout_management'))
        
    project_id = request.form.get("project_id")
    
    if not project_id:
        flash("Project ID is required", "warning")
        return redirect(url_for('layout_management'))
        
    try:
        conn = database.get_db_connection()
        c = conn.cursor()
        
        c.execute("UPDATE projects SET is_archived = 0 WHERE id = %s", (project_id,))
        conn.commit()
        conn.close()
        
        flash("Project restored successfully", "success")
        
    except Exception as e:
        app.logger.error(f"Error restoring project: {e}")
        flash("An error occurred while restoring the project", "danger")
        
    return redirect(url_for('layout_management'))


@app.route("/projects/delete", methods=["POST"])
def delete_project():
    """Delete a project"""
    if not session.get("role") == "admin":
        flash("Access denied. Admin privileges required.", "danger")
        return redirect(url_for('layout_management'))
        
    project_id = request.form.get("project_id")
    
    if not project_id:
        flash("Project ID is required", "warning")
        return redirect(url_for('layout_management'))
        
    try:
        conn = database.get_db_connection()
        c = conn.cursor()
        
        # Get project name first
        c.execute("SELECT name FROM projects WHERE id = %s", (project_id,))
        row = c.fetchone()
        if not row:
            flash("Project not found", "warning")
            conn.close()
            return redirect(url_for('layout_management'))
            
        project_name = row[0]
        
        # Check for dependencies (receipts)
        c.execute("SELECT COUNT(*) FROM receipts WHERE project_name = %s", (project_name,))
        receipt_count = c.fetchone()[0]
        
        if receipt_count > 0:
            flash(f"Cannot delete project '{project_name}'. It has {receipt_count} associated receipts.", "danger")
            conn.close()
            return redirect(url_for('layout_management'))
            
        # Delete project
        c.execute("DELETE FROM projects WHERE id = %s", (project_id,))
        conn.commit()
        conn.close()
        
        flash(f"Project '{project_name}' deleted successfully", "success")
        
    except Exception as e:
        app.logger.error(f"Error deleting project: {e}")
        flash("An error occurred while deleting the project", "danger")
        
    return redirect(url_for('layout_management'))
@app.route("/mediator_performance")
def mediator_performance():
    """Display CGM performance metrics"""
    if not (session.get("role") == "admin" or session.get("can_view_dashboard")):
        abort(403)
    
    selected_project = request.args.get("project", "").strip()
    selected_cgm = request.args.get("cgm", "").strip()
    selected_month = request.args.get("month", "").strip()  # Format: YYYY-MM or empty for all-time
    
    conn = database.get_db_connection()
    c = conn.cursor()
    
    projects = get_projects()
    
    # Get available months from commissions data
    c.execute("""
        SELECT DISTINCT DATE_FORMAT(created_at, '%Y-%m') as month
        FROM commissions
        WHERE created_at IS NOT NULL
        ORDER BY month DESC
    """)
    available_months = [row[0] for row in database.fetch_all(c)]
    
    # Get all unique CGM names for the dropdown
    if selected_project:
        c.execute("""
            SELECT DISTINCT cgm_name 
            FROM commissions 
            WHERE cgm_name IS NOT NULL 
            AND cgm_name != '' 
            AND project_name = %s
            ORDER BY cgm_name
        """, (selected_project,))
    else:
        c.execute("""
            SELECT DISTINCT cgm_name 
            FROM commissions 
            WHERE cgm_name IS NOT NULL 
            AND cgm_name != ''
            ORDER BY cgm_name
        """)
    all_cgms = [row[0] for row in database.fetch_all(c)]
    
    # Build WHERE clause based on filters
    where_clauses = []
    params = []
    
    if selected_project:
        where_clauses.append("project_name = %s")
        params.append(selected_project)
    
    if selected_cgm:
        where_clauses.append("cgm_name = %s")
        params.append(selected_cgm)
    
    if selected_month:
        where_clauses.append("DATE_FORMAT(created_at, '%Y-%m') = %s")
        params.append(selected_month)
    
    where_sql = " AND ".join(where_clauses) if where_clauses else "1=1"
    
    # Top 5 CGMs by Plots Sold (count distinct plot_no for each CGM)
    limit_sql = "" if selected_cgm else "LIMIT 5"
    c.execute(f"""
        SELECT 
            cgm_name,
            COUNT(DISTINCT plot_no) as plots_count
        FROM commissions
        WHERE cgm_name IS NOT NULL 
        AND cgm_name != ''
        AND plot_no IS NOT NULL
        AND {where_sql}
        GROUP BY cgm_name
        ORDER BY plots_count DESC
        {limit_sql}
    """, params)
    top_plots_sold = [{"cgm_name": row[0], "plots_count": row[1]} for row in database.fetch_all(c)]
    
    # Top 5 CGMs by Square Yards Sold (sum of sq_yards for each CGM)
    c.execute(f"""
        SELECT 
            cgm_name,
            SUM(sq_yards) as total_sq_yards
        FROM commissions
        WHERE cgm_name IS NOT NULL 
        AND cgm_name != ''
        AND sq_yards IS NOT NULL
        AND {where_sql}
        GROUP BY cgm_name
        ORDER BY total_sq_yards DESC
        {limit_sql}
    """, params)
    top_sq_yards = [{"cgm_name": row[0], "total_sq_yards": row[1] or 0} for row in database.fetch_all(c)]
    
    # CGM Team Total Earnings (sum of mediator_amount for each CGM's plots)
    c.execute(f"""
        SELECT 
            cgm_name,
            SUM(mediator_amount) as total_commission
        FROM commissions
        WHERE cgm_name IS NOT NULL 
        AND cgm_name != ''
        AND mediator_amount IS NOT NULL
        AND {where_sql}
        GROUP BY cgm_name
        ORDER BY total_commission DESC
        {limit_sql}
    """, params)
    top_earnings = [{"cgm_name": row[0], "total_commission": row[1] or 0} for row in database.fetch_all(c)]
    
    # Top 5 Overall Earners (across all roles)
    # Combine CGM, Sr.GM, GM, and Agent earnings
    overall_earners = []
    
    # Get CGM earnings
    c.execute(f"""
        SELECT 
            cgm_name as name,
            'CGM' as role,
            SUM(cgm_total) as total_earnings
        FROM commissions
        WHERE cgm_name IS NOT NULL 
        AND cgm_name != ''
        AND cgm_total IS NOT NULL
        AND {where_sql}
        GROUP BY cgm_name
    """, params)
    overall_earners.extend([{"name": row[0], "role": row[1], "total": row[2] or 0} for row in database.fetch_all(c)])
    
    # Get Sr.GM earnings
    c.execute(f"""
        SELECT 
            e.name,
            'Sr.GM' as role,
            SUM(e.total_amount) as total_earnings
        FROM commission_srgm_entries e
        INNER JOIN commissions c ON e.commission_id = c.id
        WHERE {where_sql.replace('cgm_name', 'c.cgm_name').replace('project_name', 'c.project_name').replace('created_at', 'c.created_at')}
        GROUP BY e.name
    """, params)
    overall_earners.extend([{"name": row[0], "role": row[1], "total": row[2] or 0} for row in database.fetch_all(c)])
    
    # Get GM earnings
    c.execute(f"""
        SELECT 
            e.name,
            'GM' as role,
            SUM(e.total_amount) as total_earnings
        FROM commission_gm_entries e
        INNER JOIN commissions c ON e.commission_id = c.id
        WHERE {where_sql.replace('cgm_name', 'c.cgm_name').replace('project_name', 'c.project_name').replace('created_at', 'c.created_at')}
        GROUP BY e.name
    """, params)
    overall_earners.extend([{"name": row[0], "role": row[1], "total": row[2] or 0} for row in database.fetch_all(c)])
    
    # Get DGM earnings
    c.execute(f"""
        SELECT 
            e.name,
            'DGM' as role,
            SUM(e.total_amount) as total_earnings
        FROM commission_dgm_entries e
        INNER JOIN commissions c ON e.commission_id = c.id
        WHERE {where_sql.replace('cgm_name', 'c.cgm_name').replace('project_name', 'c.project_name').replace('created_at', 'c.created_at')}
        GROUP BY e.name
    """, params)
    overall_earners.extend([{"name": row[0], "role": row[1], "total": row[2] or 0} for row in database.fetch_all(c)])

    # Get AGM earnings
    c.execute(f"""
        SELECT 
            e.name,
            'AGM' as role,
            SUM(e.total_amount) as total_earnings
        FROM commission_agm_entries e
        INNER JOIN commissions c ON e.commission_id = c.id
        WHERE {where_sql.replace('cgm_name', 'c.cgm_name').replace('project_name', 'c.project_name').replace('created_at', 'c.created_at')}
        GROUP BY e.name
    """, params)
    overall_earners.extend([{"name": row[0], "role": row[1], "total": row[2] or 0} for row in database.fetch_all(c)])


    
    # Sort by total earnings and get top 5
    overall_earners.sort(key=lambda x: x['total'], reverse=True)
    top_5_overall = overall_earners[:5] if not selected_cgm else overall_earners
    
    # CGM Individual Earnings (sum of cgm_total for each CGM)
    c.execute(f"""
        SELECT 
            cgm_name,
            SUM(cgm_total) as total_earnings
        FROM commissions
        WHERE cgm_name IS NOT NULL 
        AND cgm_name != ''
        AND cgm_total IS NOT NULL
        AND {where_sql}
        GROUP BY cgm_name
        ORDER BY total_earnings DESC
        {limit_sql}
    """, params)
    cgm_earnings = [{"name": row[0], "total": row[1] or 0} for row in database.fetch_all(c)]
    
    # Sr.GM Individual Earnings (from separate entries table)
    c.execute(f"""
        SELECT 
            e.name,
            SUM(e.total_amount) as total_earnings
        FROM commission_srgm_entries e
        INNER JOIN commissions c ON e.commission_id = c.id
        WHERE {where_sql.replace('cgm_name', 'c.cgm_name').replace('project_name', 'c.project_name').replace('created_at', 'c.created_at')}
        GROUP BY e.name
        ORDER BY total_earnings DESC
        {limit_sql}
    """, params)
    srgm_earnings = [{"name": row[0], "total": row[1] or 0} for row in database.fetch_all(c)]
    
    # GM Individual Earnings (from separate entries table)
    c.execute(f"""
        SELECT 
            e.name,
            SUM(e.total_amount) as total_earnings
        FROM commission_gm_entries e
        INNER JOIN commissions c ON e.commission_id = c.id
        WHERE {where_sql.replace('cgm_name', 'c.cgm_name').replace('project_name', 'c.project_name').replace('created_at', 'c.created_at')}
        GROUP BY e.name
        ORDER BY total_earnings DESC
        {limit_sql}
    """, params)
    gm_earnings = [{"name": row[0], "total": row[1] or 0} for row in database.fetch_all(c)]
    
    # DGM Individual Earnings
    c.execute(f"""
        SELECT 
            e.name,
            SUM(e.total_amount) as total_earnings
        FROM commission_dgm_entries e
        INNER JOIN commissions c ON e.commission_id = c.id
        WHERE {where_sql.replace('cgm_name', 'c.cgm_name').replace('project_name', 'c.project_name').replace('created_at', 'c.created_at')}
        GROUP BY e.name
        ORDER BY total_earnings DESC
        {limit_sql}
    """, params)
    dgm_earnings = [{"name": row[0], "total": row[1] or 0} for row in database.fetch_all(c)]

    # AGM Individual Earnings
    c.execute(f"""
        SELECT 
            e.name,
            SUM(e.total_amount) as total_earnings
        FROM commission_agm_entries e
        INNER JOIN commissions c ON e.commission_id = c.id
        WHERE {where_sql.replace('cgm_name', 'c.cgm_name').replace('project_name', 'c.project_name').replace('created_at', 'c.created_at')}
        GROUP BY e.name
        ORDER BY total_earnings DESC
        {limit_sql}
    """, params)
    agm_earnings = [{"name": row[0], "total": row[1] or 0} for row in database.fetch_all(c)]


    
    conn.close()
    
    return render_template(
        "mediator_performance.html",
        projects=projects,
        selected_project=selected_project,
        selected_cgm=selected_cgm,
        selected_month=selected_month,
        available_months=available_months,
        all_cgms=all_cgms,
        top_plots_sold=top_plots_sold,
        top_sq_yards=top_sq_yards,
        top_earnings=top_earnings,
        top_5_overall=top_5_overall,
        cgm_earnings=cgm_earnings,
        srgm_earnings=srgm_earnings,
        gm_earnings=gm_earnings,
        dgm_earnings=dgm_earnings,
        agm_earnings=agm_earnings
    )

@app.route("/mediator_details")
def mediator_details():
    """View detailed performance for a specific mediator"""
    if not (session.get("role") == "admin" or session.get("can_view_dashboard")):
        abort(403)

    name = request.args.get('name')
    role = request.args.get('role')
    project = request.args.get('project')
    month = request.args.get('month')

    if not name or not role:
        flash("Invalid request parameters", "error")
        return redirect(url_for('mediator_performance'))

    conn = database.get_db_connection()
    c = conn.cursor()

    # Build filter query
    where_clauses = ["1=1"]
    params = []

    if project:
        where_clauses.append("c.project_name = %s")
        params.append(project)
    
    if month and month != 'all':
        where_clauses.append("DATE_FORMAT(c.created_at, '%Y-%m') = %s")
        params.append(month)

    where_sql = " AND ".join(where_clauses)
    
    details = []
    total_earnings = 0
    total_sq_yards = 0
    cgm_team = "N/A"

    try:
        if role == 'CGM':
            # For CGM, they are their own team leader
            cgm_team = name
            query = f"""
                SELECT 
                    c.plot_no,
                    c.project_name,
                    c.cgm_total as earnings,
                    c.cgm_name as team_lead,
                    c.created_at,
                    c.sq_yards
                FROM commissions c
                WHERE c.cgm_name = %s
                AND c.cgm_total > 0
                AND {where_sql}
                ORDER BY c.created_at DESC
            """
            c.execute(query, [name] + params)
            
        elif role == 'Sr.GM':
            query = f"""
                SELECT 
                    c.plot_no,
                    c.project_name,
                    e.total_amount as earnings,
                    c.cgm_name as team_lead,
                    c.created_at,
                    c.sq_yards
                FROM commission_srgm_entries e
                JOIN commissions c ON e.commission_id = c.id
                WHERE e.name = %s
                AND {where_sql}
                ORDER BY c.created_at DESC
            """
            c.execute(query, [name] + params)

        elif role == 'GM':
            query = f"""
                SELECT 
                    c.plot_no,
                    c.project_name,
                    e.total_amount as earnings,
                    c.cgm_name as team_lead,
                    c.created_at,
                    c.sq_yards
                FROM commission_gm_entries e
                JOIN commissions c ON e.commission_id = c.id
                WHERE e.name = %s
                AND {where_sql}
                ORDER BY c.created_at DESC
            """
            c.execute(query, [name] + params)

        elif role == 'DGM':
            query = f"""
                SELECT 
                    c.plot_no,
                    c.project_name,
                    e.total_amount as earnings,
                    c.cgm_name as team_lead,
                    c.created_at,
                    c.sq_yards
                FROM commission_dgm_entries e
                JOIN commissions c ON e.commission_id = c.id
                WHERE e.name = %s
                AND {where_sql}
                ORDER BY c.created_at DESC
            """
            c.execute(query, [name] + params)

        elif role == 'AGM':
            query = f"""
                SELECT 
                    c.plot_no,
                    c.project_name,
                    e.total_amount as earnings,
                    c.cgm_name as team_lead,
                    c.created_at,
                    c.sq_yards
                FROM commission_agm_entries e
                JOIN commissions c ON e.commission_id = c.id
                WHERE e.name = %s
                AND {where_sql}
                ORDER BY c.created_at DESC
            """
            c.execute(query, [name] + params)



        rows = database.fetch_all(c)
        
        for row in rows:
            sq_yards = row[5] or 0
            details.append({
                'plot_no': row[0],
                'project_name': row[1],
                'earnings': row[2] or 0,
                'team_lead': row[3],
                'date': row[4],
                'sq_yards': sq_yards
            })
            total_earnings += (row[2] or 0)
            total_sq_yards += sq_yards
            
            # Update CGM team if not set (taking from the most recent transaction)
            if role != 'CGM' and row[3]:
                cgm_team = row[3]

    except Exception as e:
        print(f"Error fetching mediator details: {e}")
        flash("Error fetching details", "error")
    finally:
        conn.close()

    return render_template(
        "mediator_details.html",
        name=name,
        role=role,
        details=details,
        total_earnings=total_earnings,
        total_sq_yards=total_sq_yards,
        cgm_team=cgm_team,
        selected_project=project,
        selected_month=month
    )


@app.route("/account_summary")
def account_summary():
    if not (session.get("role") == "admin" or session.get("can_view_dashboard")):
        abort(403)
    """Display account summary with financial metrics"""
    selected_project = request.args.get("project", "").strip()
    
    conn = database.get_db_connection()
    c = conn.cursor()
    
    # Get all projects
    projects = get_projects()
    
    # Helper function to format Indian currency
    def format_indian_currency(amount):
        """Format number in Indian currency style"""
        if amount == 0:
            return "0"
        
        # Convert to string and split into integer and decimal parts
        amount_str = f"{amount:.2f}"
        parts = amount_str.split('.')
        integer_part = parts[0]
        decimal_part = parts[1] if len(parts) > 1 else "00"
        
        # Remove decimal if it's .00
        if decimal_part == "00":
            amount_str = integer_part
        
        # Apply Indian numbering system
        if len(integer_part) <= 3:
            formatted = integer_part
        else:
            # Last 3 digits
            last_three = integer_part[-3:]
            # Remaining digits
            remaining = integer_part[:-3]
            
            # Add commas every 2 digits from right to left
            formatted_remaining = ""
            for i, digit in enumerate(reversed(remaining)):
                if i > 0 and i % 2 == 0:
                    formatted_remaining = "," + formatted_remaining
                formatted_remaining = digit + formatted_remaining
            
            formatted = formatted_remaining + "," + last_three
        
        if decimal_part != "00":
            formatted += "." + decimal_part
            
        return formatted
    
    # 1. Get number of plots sold
    if selected_project:
        c.execute("""
            SELECT COUNT(DISTINCT plot_no) 
            FROM receipts 
            WHERE project_name = %s AND plot_no IS NOT NULL AND plot_no != ''
        """, (selected_project,))
    else:
        c.execute("""
            SELECT COUNT(DISTINCT plot_no) 
            FROM receipts 
            WHERE plot_no IS NOT NULL AND plot_no != ''
        """)
    
    plots_sold = database.fetch_one(c)[0] or 0
    
    # 2-4. Calculate financial metrics by aggregating plot-level data
    # We need to get unique plots and calculate their metrics
    if selected_project:
        c.execute("""
            SELECT DISTINCT plot_no 
            FROM receipts 
            WHERE project_name = %s AND plot_no IS NOT NULL AND plot_no != ''
        """, (selected_project,))
    else:
        c.execute("""
            SELECT DISTINCT plot_no, project_name
            FROM receipts 
            WHERE plot_no IS NOT NULL AND plot_no != ''
        """)
    
    plot_rows = database.fetch_all(c)
    
    total_revenue = 0
    total_paid = 0
    total_balance = 0
    
    idx_basic = get_column_index("receipts", "basic_price")
    
    for plot_row in plot_rows:
        if selected_project:
            plot_no = plot_row[0]
            project_name = selected_project
        else:
            plot_no = plot_row[0]
            project_name = plot_row[1]
        
        # Get all receipts for this plot
        c.execute("""
            SELECT * FROM receipts 
            WHERE plot_no = %s AND project_name = %s
            ORDER BY date DESC, id DESC
        """, (plot_no, project_name))
        
        rows = database.fetch_all(c)
        
        if not rows:
            continue
        
        # Get basic info from latest receipt
        latest = dict_from_row(rows[0])
        sq_yards = latest.get('square_yards', '0')
        
        try:
            sq_yards_value = float(sq_yards) if sq_yards else 0
        except (ValueError, TypeError):
            sq_yards_value = 0
        
        # Get basic price
        basic_price = 0
        if idx_basic is not None:
            for row in rows:
                if len(row) <= idx_basic:
                    continue
                raw_basic = row[idx_basic]
                if raw_basic is None or raw_basic == "":
                    continue
                try:
                    basic_price = float(raw_basic)
                    break
                except (ValueError, TypeError):
                    continue
        
        # Calculate total paid for this plot
        plot_total_paid = 0
        for row in rows:
            r = dict_from_row(row)
            plot_total_paid += (r.get("amount_numeric", 0) or 0)
        
        # If basic_price not found, use heuristic
        if (not basic_price) and sq_yards_value > 0 and plot_total_paid > 0:
            basic_price = plot_total_paid / sq_yards_value
        
        # Calculate total sale price for this plot
        plot_total_sale = basic_price * sq_yards_value
        
        # Calculate balance for this plot
        plot_balance = plot_total_sale - plot_total_paid
        
        # Add to totals
        total_revenue += plot_total_sale
        total_paid += plot_total_paid
        total_balance += plot_balance
    
    conn.close()
    
    return render_template(
        "account_summary.html",
        projects=projects,
        selected_project=selected_project,
        plots_sold=plots_sold,
        total_revenue=format_indian_currency(total_revenue),
        total_paid=format_indian_currency(total_paid),
        total_balance=format_indian_currency(total_balance)
    )


@app.route("/plots_list")
def plots_list():
    if not (session.get("role") == "admin" or session.get("can_view_dashboard")):
        abort(403)
    """Display list of all sold plots"""
    selected_project = request.args.get("project", "").strip()
    search_plot = request.args.get("search_plot", "").strip()
    
    conn = database.get_db_connection()
    c = conn.cursor()
    
    # Get all projects
    projects = get_projects()
    
    # Fetch unique plot numbers with search filter
    if selected_project and search_plot:
        # Both project and plot number specified
        c.execute("""
            SELECT DISTINCT plot_no 
            FROM receipts 
            WHERE project_name = %s AND plot_no LIKE %s AND plot_no IS NOT NULL AND plot_no != ''
            ORDER BY plot_no
        """, (selected_project, f"%{search_plot}%"))
    elif selected_project:
        # Only project specified
        c.execute("""
            SELECT DISTINCT plot_no 
            FROM receipts 
            WHERE project_name = %s AND plot_no IS NOT NULL AND plot_no != ''
            ORDER BY plot_no
        """, (selected_project,))
    elif search_plot:
        # Only plot number specified
        c.execute("""
            SELECT DISTINCT plot_no, project_name
            FROM receipts 
            WHERE plot_no LIKE %s AND plot_no IS NOT NULL AND plot_no != ''
            ORDER BY plot_no
        """, (f"%{search_plot}%",))
    else:
        # No filters - show all
        c.execute("""
            SELECT DISTINCT plot_no, project_name
            FROM receipts 
            WHERE plot_no IS NOT NULL AND plot_no != ''
            ORDER BY plot_no
        """)
    
    # If specific project: list of strings [plot1, plot2]
    # If all projects: list of tuples [(plot1, projA), (plot1, projB)]
    rows = database.fetch_all(c)
    if selected_project:
        plot_numbers = [row[0] for row in rows]
    else:
        plot_numbers = [{"plot": row[0], "project": row[1]} for row in rows]
        
    conn.close()
    
    return render_template(
        "plots_list.html",
        projects=projects,
        selected_project=selected_project,
        search_plot=search_plot,
        plots=plot_numbers,
    )


@app.route("/plot/<plot_no>")
def plot_detail(plot_no):
    if not (session.get("role") == "admin" or session.get("can_view_dashboard")):
        abort(403)
    """Display details for a specific plot"""
    conn = database.get_db_connection()
    c = conn.cursor()
    
    c.execute("""
        SELECT * FROM receipts 
        WHERE plot_no = %s 
        ORDER BY date DESC, id DESC
    """, (plot_no,))
    
    rows = database.fetch_all(c)
    conn.close()
    
    if not rows:
        abort(404)
    
    # Get the latest receipt for basic info
    latest = dict_from_row(rows[0])
    
    # Extract basic information from latest receipt
    customer_name = latest.get('customer_name', '-')
    project_name = latest.get('project_name', '-')
    sq_yards = latest.get('square_yards', '0')
    
    # Try to parse sq_yards as float
    try:
        sq_yards_value = float(sq_yards) if sq_yards else 0
    except (ValueError, TypeError):
        sq_yards_value = 0
    
    # Try to read basic price per sq yard from the DB column
    # We scan rows in order (newest first) and take the first non-empty value
    basic_price = 0
    for row in rows:
        # Access using column name directly via MySQLRow dict-like interface
        val = row.get('basic_price')
        if val is not None and val != "":
            try:
                basic_price = float(val)
                break
            except (ValueError, TypeError):
                continue

    # Let's calculate total amount paid (sum of all receipts)
    total_paid = 0
    receipts = []
    for row in rows:
        r = dict_from_row(row)
        r["amount_formatted"] = format_inr(r["amount_numeric"])
        receipts.append(r)
        total_paid += (r.get("amount_numeric", 0) or 0)

    # Calculate total sale price
    total_sale_price = basic_price * sq_yards_value
    
    # Calculate balance
    balance_amount = total_sale_price - total_paid
    
    # Format numbers in Indian format (Rs 2,32,323)
    def format_indian_currency(amount):
        """Format number in Indian currency style"""
        if amount == 0:
            return "0"
        
        # Convert to string and split into integer and decimal parts
        amount_str = f"{amount:.2f}"
        parts = amount_str.split('.')
        integer_part = parts[0]
        decimal_part = parts[1] if len(parts) > 1 else "00"
        
        # Remove decimal if it's .00
        if decimal_part == "00":
            amount_str = integer_part
        
        # Apply Indian numbering system
        if len(integer_part) <= 3:
            formatted = integer_part
        else:
            # Last 3 digits
            last_three = integer_part[-3:]
            # Remaining digits
            remaining = integer_part[:-3]
            
            # Add commas every 2 digits from right to left
            formatted_remaining = ""
            for i, digit in enumerate(reversed(remaining)):
                if i > 0 and i % 2 == 0:
                    formatted_remaining = "," + formatted_remaining
                formatted_remaining = digit + formatted_remaining
            
            formatted = formatted_remaining + "," + last_three
        
        if decimal_part != "00":
            formatted += "." + decimal_part
            
        return formatted
    
    return render_template(
        "plot_detail.html",
        plot_no=plot_no,
        project_name=project_name,
        r=latest,  # Latest receipt for basic info
        customer_name=customer_name,
        sq_yards_value=f"{sq_yards_value:.2f}" if sq_yards_value > 0 else "0",
        basic_price=format_indian_currency(basic_price),
        total_sale=format_indian_currency(total_sale_price),
        amount_paid=format_indian_currency(total_paid),
        balance=format_indian_currency(balance_amount),
        history=receipts
    )


# -----------------------------
# Commission Calculator
# -----------------------------
@app.route("/commission_calculator", methods=["GET", "POST"])
def commission_calculator():
    """Commission calculator with PDF generation"""
    if request.method == "GET":
        projects = get_projects()
        return render_template("commission_calculator.html", projects=projects)



    # POST: Generate PDF
    try:
        # Extract form data
        # Handle multiple Sr. GM, GM, and Agent entries
        srgm_names = request.form.getlist('srgm_name[]')
        srgm_rates = request.form.getlist('srgm_rate[]')
        gm_names = request.form.getlist('gm_name[]')
        gm_rates = request.form.getlist('gm_rate[]')
        dgm_names = request.form.getlist('dgm_name[]')
        dgm_rates = request.form.getlist('dgm_rate[]')
        agm_names = request.form.getlist('agm_name[]')
        agm_rates = request.form.getlist('agm_rate[]')

        
        # Helper to safely parse float with commas (defined early for use here)
        def safe_float_early(val):
            if not val:
                return 0.0
            if isinstance(val, (int, float)):
                return float(val)
            try:
                return float(str(val).replace(',', '').strip())
            except (ValueError, TypeError):
                return 0.0

        # Join names with commas and sum rates (for database storage)
        srgm_name_combined = ', '.join([n.strip() for n in srgm_names if n.strip()])
        srgm_rate_total = sum([safe_float_early(r) for r in srgm_rates if r])
        gm_name_combined = ', '.join([n.strip() for n in gm_names if n.strip()])
        gm_rate_total = sum([safe_float_early(r) for r in gm_rates if r])
        dgm_name_combined = ', '.join([n.strip() for n in dgm_names if n.strip()])
        dgm_rate_total = sum([safe_float_early(r) for r in dgm_rates if r])
        agm_name_combined = ', '.join([n.strip() for n in agm_names if n.strip()])
        agm_rate_total = sum([safe_float_early(r) for r in agm_rates if r])

        
        # Store individual entries for PDF generation
        srgm_entries = [(n.strip(), safe_float_early(r)) for n, r in zip(srgm_names, srgm_rates) if n.strip() or (r and safe_float_early(r) > 0)]
        gm_entries = [(n.strip(), safe_float_early(r)) for n, r in zip(gm_names, gm_rates) if n.strip() or (r and safe_float_early(r) > 0)]
        dgm_entries = [(n.strip(), safe_float_early(r)) for n, r in zip(dgm_names, dgm_rates) if n.strip() or (r and safe_float_early(r) > 0)]
        agm_entries = [(n.strip(), safe_float_early(r)) for n, r in zip(agm_names, agm_rates) if n.strip() or (r and safe_float_early(r) > 0)]

        
        app.logger.info(f"DGM Entries: {dgm_entries}")
        app.logger.info(f"AGM Entries: {agm_entries}")
        
        # Create breakdown dict for database storage
        commission_breakdown = {
            'srgm': srgm_entries,
            'gm': gm_entries,
            'dgm': dgm_entries,
            'agm': agm_entries,

        }
        
        # Helper to safely parse float with commas
        def safe_float(val):
            if not val:
                return 0.0
            if isinstance(val, (int, float)):
                return float(val)
            try:
                return float(str(val).replace(',', '').strip())
            except (ValueError, TypeError):
                return 0.0

        # Form data extraction with safe float parsing
        form_data = {
            'plot_no': request.form.get('plot_no', '').strip(),
            'project_name': request.form.get('project_name', '').strip(),
            'sq_yards': safe_float(request.form.get('sq_yards', 0)),
            'original_price': safe_float(request.form.get('original_price', 0)),
            'negotiated_price': safe_float(request.form.get('negotiated_price', 0)),
            'advance_received': safe_float(request.form.get('advance_received', 0)),

            'agreement_percentage': safe_float(request.form.get('agreement_percentage', 0)) / 100,
            'amount_paid_at_agreement': safe_float(request.form.get('amount_paid_at_agreement', 0)),
            'amc_charges': safe_float(request.form.get('amc_charges', 0)),
            'mediator_deduction': safe_float(request.form.get('mediator_deduction', 0)),
            'broker_commission': safe_float(request.form.get('broker_commission', 0)),
            'cgm_rate': safe_float(request.form.get('cgm_rate', 0)),
            'srgm_rate': srgm_rate_total,
            'gm_rate': gm_rate_total,
            'dgm_rate': dgm_rate_total,
            'agm_rate': agm_rate_total,

            'cgm_name': request.form.get('cgm_name', '').strip(),
            'srgm_name': srgm_name_combined,
            'gm_name': gm_name_combined,
            'dgm_name': dgm_name_combined,
            'agm_name': agm_name_combined,

            'commission_breakdown': json.dumps(commission_breakdown),
            # Individual entries for PDF
            'srgm_entries': srgm_entries,
            'gm_entries': gm_entries,
            'dgm_entries': dgm_entries,
            'agm_entries': agm_entries,

        }
        
        # Calculate all values (same logic as JavaScript)
        calculations = calculate_commission(form_data)
        
        # Check if we're editing or creating new
        commission_id_str = request.form.get('commission_id', '').strip()
        
        if commission_id_str:
            # UPDATE existing commission
            commission_id = int(commission_id_str)
            update_commission_in_db(commission_id, form_data, calculations)
        else:
            # INSERT new commission
            commission_id = save_commission_to_db(form_data, calculations)
        

        
        # Create filename with plot number
        plot_no_safe = _safe_filename_fragment(form_data['plot_no'], 'commission')
        filename_pdf = f"commission_Plot{plot_no_safe}.pdf"
        filename_docx = f"commission_Plot{plot_no_safe}.docx"
        
        app.logger.info(f"Saved commission for plot {filename_pdf} with ID {commission_id}")
        
        # Return JSON response with preview URL and docx URL
        return jsonify({
            'success': True,
            'preview_url': url_for('preview_commission_pdf', commission_id=commission_id),
            'docx_url': url_for('download_commission_docx', commission_id=commission_id, filename=filename_docx),
            'docx_filename': filename_docx
        })
        
    except Exception as e:
        app.logger.exception("Error generating commission PDF")
        return jsonify({
            'success': False,
            'error': f"Error generating PDF: {str(e)}"
        }), 500


@app.route("/commission/edit/<int:commission_id>")
def edit_commission(commission_id):
    """Edit existing commission calculation"""
    conn = database.get_db_connection()
    c = conn.cursor()
    
    c.execute("SELECT * FROM commissions WHERE id = %s", (commission_id,))
    row = database.fetch_one(c)
    conn.close()
    
    if not row:
        flash("Commission not found", "danger")
        return redirect(url_for('search_commission'))
    
    # Convert row to dict for template
    commission_data = dict(row)
    
    # Try to load breakdown from JSON
    breakdown = None
    if commission_data.get('commission_breakdown'):
        try:
            breakdown = json.loads(commission_data['commission_breakdown'])
        except:
            pass
            
    if breakdown:
        # Use stored breakdown
        # Convert lists of [name, rate] to dicts {'name': name, 'rate': rate} for template
        commission_data['srgm_entries'] = [{'name': e[0], 'rate': e[1]} for e in breakdown.get('srgm', [])]
        commission_data['gm_entries'] = [{'name': e[0], 'rate': e[1]} for e in breakdown.get('gm', [])]
        commission_data['agent_entries'] = [{'name': e[0], 'rate': e[1]} for e in breakdown.get('agent', [])]
    else:
        # Fallback: Parse combined names into individual entries
        # SrGM
        srgm_names = [n.strip() for n in (commission_data.get('srgm_name') or '').split(',') if n.strip()]
        srgm_total_rate = commission_data.get('srgm_rate') or 0
        commission_data['srgm_entries'] = []
        if srgm_names:
            # Distribute rate equally among all names
            rate_per_person = srgm_total_rate / len(srgm_names)
            for name in srgm_names:
                commission_data['srgm_entries'].append({'name': name, 'rate': rate_per_person})
        elif srgm_total_rate > 0:
             # If no name but rate exists, create empty name entry
             commission_data['srgm_entries'].append({'name': '', 'rate': srgm_total_rate})

        # GM
        gm_names = [n.strip() for n in (commission_data.get('gm_name') or '').split(',') if n.strip()]
        gm_total_rate = commission_data.get('gm_rate') or 0
        commission_data['gm_entries'] = []
        if gm_names:
            rate_per_person = gm_total_rate / len(gm_names)
            for name in gm_names:
                commission_data['gm_entries'].append({'name': name, 'rate': rate_per_person})
        elif gm_total_rate > 0:
             commission_data['gm_entries'].append({'name': '', 'rate': gm_total_rate})

        # DGM
        dgm_names = [n.strip() for n in (commission_data.get('dgm_name') or '').split(',') if n.strip()]
        dgm_total_rate = commission_data.get('dgm_rate') or 0
        commission_data['dgm_entries'] = []
        if dgm_names:
            rate_per_person = dgm_total_rate / len(dgm_names)
            for name in dgm_names:
                commission_data['dgm_entries'].append({'name': name, 'rate': rate_per_person})
        elif dgm_total_rate > 0:
             commission_data['dgm_entries'].append({'name': '', 'rate': dgm_total_rate})

        # AGM
        agm_names = [n.strip() for n in (commission_data.get('agm_name') or '').split(',') if n.strip()]
        agm_total_rate = commission_data.get('agm_rate') or 0
        commission_data['agm_entries'] = []
        if agm_names:
            rate_per_person = agm_total_rate / len(agm_names)
            for name in agm_names:
                commission_data['agm_entries'].append({'name': name, 'rate': rate_per_person})
        elif agm_total_rate > 0:
             commission_data['agm_entries'].append({'name': '', 'rate': agm_total_rate})


            
    # Load projects for dropdown (same as in commission_calculator GET)
    projects = get_projects()

    app.logger.info(f"Editing commission {commission_id}: {commission_data}")
    
    return render_template(
        "commission_calculator.html",
        edit_mode=True,
        commission_id=commission_id,
        commission_data=commission_data,
        projects=projects,
    )


@app.route("/commission/preview/<int:commission_id>")
def preview_commission_pdf(commission_id):
    """Render PDF preview wrapper page"""
    conn = database.get_db_connection()
    c = conn.cursor()
    c.execute("SELECT plot_no FROM commissions WHERE id = %s", (commission_id,))
    row = database.fetch_one(c)
    conn.close()
    
    if not row:
        abort(404)
        
    plot_no = row[0]
    plot_no_safe = _safe_filename_fragment(plot_no, 'commission')
    filename = f"commission_Plot{plot_no_safe}.pdf"
    
    return render_template(
        "pdf_preview.html",
        filename=filename,
        raw_url=url_for('raw_commission_pdf', commission_id=commission_id, filename=filename),
        commission_id=commission_id
    )


@app.route("/commission/raw/<int:commission_id>/<filename>")
def raw_commission_pdf(commission_id, filename):
    """View commission PDF by ID"""
    conn = database.get_db_connection()
    c = conn.cursor()
    
    c.execute("SELECT * FROM commissions WHERE id = %s", (commission_id,))
    row = database.fetch_one(c)
    conn.close()
    
    if not row:
        abort(404)
    
    # row is already a MySQLRow (dict-like object), so we can use it directly
    commission_data = dict(row)
    
    # Create form_data dict for PDF generation
    form_data = {
        'plot_no': commission_data.get('plot_no', ''),
        'project_name': commission_data.get('project_name', ''),
        'sq_yards': commission_data.get('sq_yards', 0),
        'original_price': commission_data.get('original_price', 0),
        'negotiated_price': commission_data.get('negotiated_price', 0),
        'advance_received': commission_data.get('advance_received', 0),
        'agent_commission_rate': commission_data.get('agent_commission_rate', 0),
        'agreement_percentage': commission_data.get('agreement_percentage', 0),
        'amount_paid_at_agreement': commission_data.get('amount_paid_at_agreement', 0),
        'amc_charges': commission_data.get('amc_charges', 0),
        'mediator_deduction': commission_data.get('mediator_deduction', 0),
        'broker_commission': commission_data.get('broker_commission', 0),
        'cgm_rate': commission_data.get('cgm_rate', 0),
        'srgm_rate': commission_data.get('srgm_rate', 0),
        'gm_rate': commission_data.get('gm_rate', 0),
        'dgm_rate': commission_data.get('dgm_rate', 0),
        'agm_rate': commission_data.get('agm_rate', 0),
        'agent_rate': commission_data.get('agent_rate', 0),
        'cgm_name': commission_data.get('cgm_name', ''),
        'srgm_name': commission_data.get('srgm_name', ''),
        'gm_name': commission_data.get('gm_name', ''),
        'dgm_name': commission_data.get('dgm_name', ''),
        'agm_name': commission_data.get('agm_name', ''),
        'agent_name': commission_data.get('agent_name', ''),
        'commission_breakdown': commission_data.get('commission_breakdown', '')
    }
    
    # Parse commission_breakdown JSON to reconstruct individual entries
    breakdown = None
    if commission_data.get('commission_breakdown'):
        try:
            breakdown = json.loads(commission_data['commission_breakdown'])
        except:
            pass
    
    if breakdown:
        # Reconstruct individual entries for PDF generation
        form_data['srgm_entries'] = [(e[0], e[1]) for e in breakdown.get('srgm', [])]
        form_data['gm_entries'] = [(e[0], e[1]) for e in breakdown.get('gm', [])]
        form_data['dgm_entries'] = [(e[0], e[1]) for e in breakdown.get('dgm', [])]
        form_data['agm_entries'] = [(e[0], e[1]) for e in breakdown.get('agm', [])]

    else:
        # Fallback: create single entries from combined names and total rates
        if commission_data.get('srgm_name'):
            srgm_names = [name.strip() for name in commission_data['srgm_name'].split(',') if name.strip()]
            srgm_rate = float(commission_data.get('srgm_rate', 0) or 0)
            srgm_rate_per = srgm_rate / len(srgm_names) if srgm_names else 0
            form_data['srgm_entries'] = [(name, srgm_rate_per) for name in srgm_names]
        else:
            form_data['srgm_entries'] = []
            
        if commission_data.get('gm_name'):
            gm_names = [name.strip() for name in commission_data['gm_name'].split(',') if name.strip()]
            gm_rate = float(commission_data.get('gm_rate', 0) or 0)
            gm_rate_per = gm_rate / len(gm_names) if gm_names else 0
            form_data['gm_entries'] = [(name, gm_rate_per) for name in gm_names]
        else:
            form_data['gm_entries'] = []
            
        if commission_data.get('dgm_name'):
            dgm_names = [name.strip() for name in commission_data['dgm_name'].split(',') if name.strip()]
            dgm_rate = float(commission_data.get('dgm_rate', 0) or 0)
            dgm_rate_per = dgm_rate / len(dgm_names) if dgm_names else 0
            form_data['dgm_entries'] = [(name, dgm_rate_per) for name in dgm_names]
        else:
            form_data['dgm_entries'] = []

        if commission_data.get('agm_name'):
            agm_names = [name.strip() for name in commission_data['agm_name'].split(',') if name.strip()]
            agm_rate = float(commission_data.get('agm_rate', 0) or 0)
            agm_rate_per = agm_rate / len(agm_names) if agm_names else 0
            form_data['agm_entries'] = [(name, agm_rate_per) for name in agm_names]
        else:
            form_data['agm_entries'] = []
            

    
    # Create calculations dict
    calculations = {
        'total_amount': commission_data.get('total_amount', 0),
        'w_value': commission_data.get('w_value', 0),
        'b_value': commission_data.get('b_value', 0),
        'balance_amount': commission_data.get('balance_amount', 0),
        'actual_agreement_amount': commission_data.get('actual_agreement_amount', 0),
        'agreement_balance': commission_data.get('agreement_balance', 0),
        'mediator_amount': commission_data.get('mediator_amount', 0),
        'mediator_deduction': commission_data.get('mediator_deduction', 0),
        'mediator_actual_payment': commission_data.get('mediator_actual_payment', 0),
        'mediator_at_agreement': commission_data.get('mediator_at_agreement', 0),
        'cgm_total': commission_data.get('cgm_total', 0),
        'cgm_at_agreement': commission_data.get('cgm_at_agreement', 0),
        'cgm_at_registration': commission_data.get('cgm_at_registration', 0),
        'srgm_total': commission_data.get('srgm_total', 0),
        'srgm_at_agreement': commission_data.get('srgm_at_agreement', 0),
        'srgm_at_registration': commission_data.get('srgm_at_registration', 0),
        'gm_total': commission_data.get('gm_total', 0),
        'gm_at_agreement': commission_data.get('gm_at_agreement', 0),
        'gm_at_registration': commission_data.get('gm_at_registration', 0),

    }
    
    # Generate PDF using existing function
    pdf_bytes = generate_commission_pdf_bytes(form_data, calculations)
    
    return send_file(
        io.BytesIO(pdf_bytes),
        as_attachment=False,
        download_name=filename,
        mimetype='application/pdf'
    )


@app.route("/users", methods=["GET", "POST"])
def user_management():
    if session.get("role") != "admin":
        abort(403)

    conn = database.get_db_connection()
    c = conn.cursor()

    error = None
    message = None

    if request.method == "POST":
        action = request.form.get("action", "").strip()

        if action == "create":
            username = (request.form.get("username") or "").strip()
            password = (request.form.get("password") or "").strip()
            role = (request.form.get("role") or "user").strip() or "user"
            can_view_dashboard = 1 if request.form.get("can_view_dashboard") == "on" else 0
            can_search_receipts = 1 if request.form.get("can_search_receipts") == "on" else 0
            can_view_vishvam_layout = 1 if request.form.get("can_view_vishvam_layout") == "on" else 0

            if not username or not password:
                error = "Username and password are required to create a user."
            else:
                created_at = datetime.utcnow().isoformat()
                # Use pbkdf2:sha256 explicitly to avoid hashlib.scrypt requirement
                pwd_hash = generate_password_hash(password, method="pbkdf2:sha256")
                try:
                    c.execute(
                        "INSERT INTO users (username, password_hash, role, can_view_dashboard, can_search_receipts, can_view_vishvam_layout, created_at) VALUES (%s, %s, %s, %s, %s, %s, %s)",
                        (username, pwd_hash, role, can_view_dashboard, can_search_receipts, can_view_vishvam_layout, created_at),
                    )
                    conn.commit()
                    message = f"User '{username}' created."
                except database.IntegrityError:
                    error = "Username already exists. Choose a different one."

        elif action == "update":
            user_id = request.form.get("user_id")
            role = (request.form.get("role") or "user").strip() or "user"
            can_view_dashboard = 1 if request.form.get("can_view_dashboard") == "on" else 0
            can_search_receipts = 1 if request.form.get("can_search_receipts") == "on" else 0
            can_view_vishvam_layout = 1 if request.form.get("can_view_vishvam_layout") == "on" else 0

            if user_id:
                try:
                    # Prevent removing own admin role accidentally: only apply if updating another user
                    if int(user_id) == session.get("user_id") and role != "admin":
                        error = "You cannot remove your own admin role."
                    else:
                        c.execute(
                            "UPDATE users SET role = %s, can_view_dashboard = %s, can_search_receipts = %s, can_view_vishvam_layout = %s WHERE id = %s",
                            (role, can_view_dashboard, can_search_receipts, can_view_vishvam_layout, user_id),
                        )
                        conn.commit()
                        message = "User updated."
                except Exception as e:
                    error = f"Error updating user: {e}"

        elif action == "delete":
            user_id = request.form.get("user_id")
            if user_id:
                try:
                    uid = int(user_id)
                    # Prevent deleting own account
                    if uid == session.get("user_id"):
                        error = "You cannot delete your own account while logged in."
                    else:
                        # Check if this user is an admin and whether they are the last admin
                        c.execute("SELECT role FROM users WHERE id = %s", (uid,))
                        row = database.fetch_one(c)
                        if not row:
                            error = "User not found."
                        else:
                            user_role = row[0]
                            if user_role == "admin":
                                c.execute("SELECT COUNT(*) FROM users WHERE role = 'admin'")
                                admin_count = database.fetch_one(c)[0] or 0
                                if admin_count <= 1:
                                    error = "You cannot delete the last remaining admin user."
                                else:
                                    c.execute("DELETE FROM users WHERE id = %s", (uid,))
                                    conn.commit()
                                    message = "User deleted."
                            else:
                                c.execute("DELETE FROM users WHERE id = %s", (uid,))
                                conn.commit()
                                message = "User deleted."
                except Exception as e:
                    error = f"Error deleting user: {e}"

    # Fetch all users
    c.execute("SELECT id, username, role, can_view_dashboard, can_search_receipts, can_view_vishvam_layout, created_at FROM users ORDER BY created_at DESC")
    users = [dict(row) for row in database.fetch_all(c)]
    conn.close()

    if error:
        flash(error, "danger")
    if message:
        flash(message, "success")

    return render_template("user_management.html", users=users)


@app.route("/pending_receipts")
def pending_receipts():
    """Admin-only: View all pending receipts awaiting approval"""
    if session.get("role") != "admin":
        abort(403)
    
    conn = database.get_db_connection()
    c = conn.cursor()
    
    c.execute("""
        SELECT * FROM pending_receipts 
        WHERE status = 'pending'
        ORDER BY submitted_at DESC
    """)
    
    pending = [dict(row) for row in database.fetch_all(c)]
    conn.close()
    
    return render_template("pending_receipts.html", pending_receipts=pending)


@app.route("/approve_receipt/<int:pending_id>", methods=["GET", "POST"])
def approve_receipt(pending_id):
    """Admin-only: Review and approve/edit a pending receipt"""
    if session.get("role") != "admin":
        abort(403)
    
    conn = database.get_db_connection()
    c = conn.cursor()
    
    if request.method == "GET":
        # Load pending receipt for review
        c.execute("SELECT * FROM pending_receipts WHERE id = %s", (pending_id,))
        pending = database.fetch_one(c)
        conn.close()
        
        if not pending:
            flash("Pending receipt not found", "danger")
            return redirect(url_for("pending_receipts"))
        
        # Get projects for dropdown
        projects = get_projects()
        return render_template("approve_receipt.html", receipt=dict(pending), projects=projects, pending_id=pending_id)
    
    else:  # POST - Approve or Reject
        action = request.form.get("action", "").strip()
        
        if action == "approve":
            # Get (possibly edited) data from form
            form = request.form
            no = form.get("no", "").strip()
            project_name = form.get("project_name", "").strip()
            date = form.get("date", "").strip()
            venture = form.get("venture", "").strip()
            customer_name = form.get("customer_name", "").strip()
            amount_numeric = form.get("amount_numeric", "0").replace(",", "").strip()
            amount_words = form.get("amount_words", "").strip() or number_to_words(amount_numeric)
            plot_no = form.get("plot_no", "").strip()
            square_yards = form.get("square_yards", "").strip()
            purpose = form.get("purpose", "").strip()
            drawn_bank = form.get("drawn_bank", "").strip()
            branch = form.get("branch", "").strip()
            payment_mode = form.get("payment_mode", "").strip()
            instrument_no = form.get("instrument_no", "").strip()
            
            created_at = datetime.utcnow().isoformat()
            
            # Insert into main receipts table
            c.execute(
                """
                INSERT INTO receipts
                (no, project_name, date, venture, customer_name, amount_numeric, amount_words,
                 plot_no, square_yards, purpose, drawn_bank, branch, payment_mode, instrument_no, created_at)
                VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
            """,
                (
                    no, project_name, date, venture, customer_name, amount_numeric, amount_words,
                    plot_no, square_yards, purpose, drawn_bank, branch, payment_mode, instrument_no, created_at
                ),
            )
            receipt_id = c.lastrowid
            
            # Mark pending receipt as approved and delete
            c.execute("DELETE FROM pending_receipts WHERE id = %s", (pending_id,))
            conn.commit()
            conn.close()
            
            flash("Receipt approved and saved successfully!", "success")
            return redirect(url_for("view_receipt", receipt_id=receipt_id))
        
        elif action == "reject":
            # Delete the pending receipt
            c.execute("DELETE FROM pending_receipts WHERE id = %s", (pending_id,))
            conn.commit()
            conn.close()
            
            flash("Receipt rejected and removed from pending list.", "warning")
            return redirect(url_for("pending_receipts"))
        
        conn.close()
        return redirect(url_for("pending_receipts"))

    conn = database.get_db_connection()
    c = conn.cursor()
    
    c.execute("SELECT * FROM commissions WHERE id = %s", (commission_id,))
    row = database.fetch_one(c)
    conn.close()
    
    if not row:
        abort(404)
        
    # Convert row to dict to use .get() method
    row = dict(row)
        
    # Reconstruct form_data and calculations from DB row
    form_data = {
        'plot_no': row['plot_no'],
        'sq_yards': row['sq_yards'],
        'original_price': row['original_price'],
        'negotiated_price': row['negotiated_price'],
        'advance_received': row['advance_received'],

        'agreement_percentage': row['agreement_percentage'],
        'amount_paid_at_agreement': row['amount_paid_at_agreement'],
        'amc_charges': row['amc_charges'],
        'cgm_rate': row['cgm_rate'],
        'srgm_rate': row['srgm_rate'],
        'gm_rate': row['gm_rate'],
        'dgm_rate': row['dgm_rate'],
        'agm_rate': row['agm_rate'],

        'cgm_name': row.get('cgm_name', ''),
        'srgm_name': row.get('srgm_name', ''),
        'gm_name': row.get('gm_name', ''),
        'dgm_name': row.get('dgm_name', ''),
        'agm_name': row.get('agm_name', ''),

    }
    
    calculations = {
        'total_amount': row['total_amount'],
        'w_value': row['w_value'],
        'b_value': row['b_value'],
        'balance_amount': row['balance_amount'],
        'actual_agreement_amount': row['actual_agreement_amount'],
        'agreement_balance': row['agreement_balance'],

        'cgm_total': row['cgm_total'],
        'cgm_at_agreement': row['cgm_at_agreement'],
        'cgm_at_registration': row['cgm_at_registration'],
        'srgm_total': row['srgm_total'],
        'srgm_at_agreement': row['srgm_at_agreement'],
        'srgm_at_registration': row['srgm_at_registration'],
        'gm_total': row['gm_total'],
        'gm_at_agreement': row['gm_at_agreement'],
        'gm_at_registration': row['gm_at_registration'],

    }
    
    # Try to load breakdown from JSON
    breakdown = None
    if row.get('commission_breakdown'):
        try:
            breakdown = json.loads(row['commission_breakdown'])
        except:
            pass
            
    if breakdown:
        # Use stored breakdown for PDF
        # Ensure entries are tuples (name, rate) as expected by PDF generator
        form_data['srgm_entries'] = [(e[0], e[1]) for e in breakdown.get('srgm', [])]
        form_data['gm_entries'] = [(e[0], e[1]) for e in breakdown.get('gm', [])]
        form_data['dgm_entries'] = [(e[0], e[1]) for e in breakdown.get('dgm', [])]
        form_data['agm_entries'] = [(e[0], e[1]) for e in breakdown.get('agm', [])]

    else:
        # Fallback: Parse combined names into individual entries
        # SrGM
        srgm_names = [n.strip() for n in (form_data.get('srgm_name') or '').split(',') if n.strip()]
        srgm_total_rate = form_data.get('srgm_rate') or 0
        form_data['srgm_entries'] = []
        if srgm_names:
            # Distribute rate equally among all names
            rate_per_person = srgm_total_rate / len(srgm_names)
            for name in srgm_names:
                form_data['srgm_entries'].append((name, rate_per_person))
        elif srgm_total_rate > 0:
             form_data['srgm_entries'].append(('', srgm_total_rate))

        # GM
        gm_names = [n.strip() for n in (form_data.get('gm_name') or '').split(',') if n.strip()]
        gm_total_rate = form_data.get('gm_rate') or 0
        form_data['gm_entries'] = []
        if gm_names:
            rate_per_person = gm_total_rate / len(gm_names)
            for name in gm_names:
                form_data['gm_entries'].append((name, rate_per_person))
        elif gm_total_rate > 0:
             form_data['gm_entries'].append(('', gm_total_rate))

        # DGM
        dgm_names = [n.strip() for n in (form_data.get('dgm_name') or '').split(',') if n.strip()]
        dgm_total_rate = form_data.get('dgm_rate') or 0
        form_data['dgm_entries'] = []
        if dgm_names:
            rate_per_person = dgm_total_rate / len(dgm_names)
            for name in dgm_names:
                form_data['dgm_entries'].append((name, rate_per_person))
        elif dgm_total_rate > 0:
             form_data['dgm_entries'].append(('', dgm_total_rate))

        # AGM
        agm_names = [n.strip() for n in (form_data.get('agm_name') or '').split(',') if n.strip()]
        agm_total_rate = form_data.get('agm_rate') or 0
        form_data['agm_entries'] = []
        if agm_names:
            rate_per_person = agm_total_rate / len(agm_names)
            for name in agm_names:
                form_data['agm_entries'].append((name, rate_per_person))
        elif agm_total_rate > 0:
             form_data['agm_entries'].append(('', agm_total_rate))


    
    # Generate PDF
    pdf_bytes = generate_commission_pdf_bytes(form_data, calculations)
    
    # Create filename
    plot_no_safe = _safe_filename_fragment(form_data['plot_no'], 'commission')
    filename = f"commission_Plot{plot_no_safe}.pdf"
    
    # Return PDF inline for preview
    response = make_response(pdf_bytes)
    response.headers['Content-Type'] = 'application/pdf'
    # Inline disposition allows browser preview
    response.headers['Content-Disposition'] = f'inline; filename="{filename}"'
    
    return response


def calculate_commission(form_data):
    """Calculate all commission values based on input data"""
    sq_yards = form_data['sq_yards']
    negotiated_price = form_data['negotiated_price']
    original_price = form_data['original_price']
    advance_received = form_data['advance_received']

    agreement_percentage = form_data['agreement_percentage']
    amount_paid_at_agreement = form_data['amount_paid_at_agreement']
    amc_charges = form_data['amc_charges']
    
    # Basic calculations
    total_amount = (sq_yards * negotiated_price) + (amc_charges * sq_yards)
    w_value = sq_yards * 5500
    b_value = total_amount - w_value
    balance_amount = total_amount - advance_received
    actual_agreement_amount = total_amount * agreement_percentage
    agreement_balance = actual_agreement_amount - amount_paid_at_agreement - advance_received
    

    
    # CGM calculations
    cgm_total = form_data['cgm_rate'] * sq_yards
    cgm_at_agreement = cgm_total * agreement_percentage
    cgm_at_registration = cgm_total - cgm_at_agreement
    
    # SrGM calculations
    srgm_total = form_data['srgm_rate'] * sq_yards
    srgm_at_agreement = srgm_total * agreement_percentage
    srgm_at_registration = srgm_total - srgm_at_agreement
    
    # GM calculations
    gm_total = form_data['gm_rate'] * sq_yards
    gm_at_agreement = gm_total * agreement_percentage
    gm_at_registration = gm_total - gm_at_agreement

    # DGM calculations
    dgm_total = form_data['dgm_rate'] * sq_yards
    dgm_at_agreement = dgm_total * agreement_percentage
    dgm_at_registration = dgm_total - dgm_at_agreement

    # AGM calculations
    agm_total = form_data['agm_rate'] * sq_yards
    agm_at_agreement = agm_total * agreement_percentage
    agm_at_registration = agm_total - agm_at_agreement
    

    
    return {
        'total_amount': total_amount,
        'w_value': w_value,
        'b_value': b_value,
        'balance_amount': balance_amount,
        'actual_agreement_amount': actual_agreement_amount,
        'agreement_balance': agreement_balance,

        'cgm_total': cgm_total,
        'cgm_at_agreement': cgm_at_agreement,
        'cgm_at_registration': cgm_at_registration,
        'srgm_total': srgm_total,
        'srgm_at_agreement': srgm_at_agreement,
        'srgm_at_registration': srgm_at_registration,
        'gm_total': gm_total,
        'gm_at_agreement': gm_at_agreement,
        'gm_at_registration': gm_at_registration,
        'dgm_total': dgm_total,
        'dgm_at_agreement': dgm_at_agreement,
        'dgm_at_registration': dgm_at_registration,
        'agm_total': agm_total,
        'agm_at_agreement': agm_at_agreement,
        'agm_at_registration': agm_at_registration,

    }


def save_commission_to_db(form_data, calculations):
    """Save commission calculation to database"""
    conn = database.get_db_connection()
    c = conn.cursor()
    
    c.execute("""
        INSERT INTO commissions (
            plot_no, project_name, sq_yards, original_price, negotiated_price,
            advance_received,
            agreement_percentage,
            amount_paid_at_agreement, amc_charges, mediator_deduction, cgm_rate, srgm_rate,
            gm_rate, dgm_rate, agm_rate, 
            cgm_name, srgm_name, gm_name, dgm_name, agm_name,
            amount_paid_at_agreement, amc_charges, mediator_deduction, broker_commission, cgm_rate, srgm_rate,
            gm_rate, dgm_rate, agm_rate, agent_rate,
            cgm_name, srgm_name, gm_name, dgm_name, agm_name, agent_name,
            total_amount, w_value, b_value,
            balance_amount, actual_agreement_amount, agreement_balance,
            cgm_total, cgm_at_agreement,
            cgm_at_registration, srgm_total, srgm_at_agreement,
            srgm_at_registration, gm_total, gm_at_agreement,
            gm_at_registration, dgm_total, dgm_at_agreement, dgm_at_registration,
            agm_total, agm_at_agreement, agm_at_registration,
            created_by, commission_breakdown
        ) VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s,
                  %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
    """, (
        form_data['plot_no'], form_data.get('project_name', ''), form_data['sq_yards'], form_data['original_price'],
        form_data['negotiated_price'], form_data['advance_received'],
        form_data['agreement_percentage'],
        form_data['amount_paid_at_agreement'], form_data['amc_charges'], form_data.get('mediator_deduction', 0),
        form_data.get('broker_commission', 0),
        form_data['cgm_rate'], form_data['srgm_rate'], form_data['gm_rate'],
        form_data['dgm_rate'], form_data['agm_rate'], form_data.get('agent_rate', 0),
        form_data.get('cgm_name', ''), form_data.get('srgm_name', ''),
        form_data.get('gm_name', ''), form_data.get('dgm_name', ''), form_data.get('agm_name', ''), form_data.get('agent_name', ''),
        calculations['total_amount'], calculations['w_value'],
        calculations['b_value'], calculations['balance_amount'],
        calculations['actual_agreement_amount'], calculations['agreement_balance'],
        calculations['cgm_total'], calculations['cgm_at_agreement'],
        calculations['cgm_at_registration'], calculations['srgm_total'],
        calculations['srgm_at_agreement'], calculations['srgm_at_registration'],
        calculations['gm_total'], calculations['gm_at_agreement'],
        calculations['gm_at_registration'], 
        calculations['dgm_total'], calculations['dgm_at_agreement'], calculations['dgm_at_registration'],
        calculations['agm_total'], calculations['agm_at_agreement'], calculations['agm_at_registration'],
        session.get('username', 'admin'),
        form_data.get('commission_breakdown', '')
    ))
    
    commission_id = c.lastrowid
    
    # Save individual Sr.GM entries
    srgm_entries = form_data.get('srgm_entries', [])
    if srgm_entries:
        for entry in srgm_entries:
            name = entry[0] if isinstance(entry, (list, tuple)) else entry.get('name', '')
            rate = entry[1] if isinstance(entry, (list, tuple)) and len(entry) > 1 else entry.get('rate', 0)
            
            if name and rate:
                sq_yards = form_data['sq_yards']
                agreement_pct = form_data['agreement_percentage']
                
                total = rate * sq_yards
                agreement = total * agreement_pct
                registration = total - agreement
                
                c.execute("""
                    INSERT INTO commission_srgm_entries 
                    (commission_id, name, total_amount, at_agreement, at_registration)
                    VALUES (%s, %s, %s, %s, %s)
                """, (commission_id, name.strip(), total, agreement, registration))
    
    # Save individual GM entries
    gm_entries = form_data.get('gm_entries', [])
    if gm_entries:
        for entry in gm_entries:
            name = entry[0] if isinstance(entry, (list, tuple)) else entry.get('name', '')
            rate = entry[1] if isinstance(entry, (list, tuple)) and len(entry) > 1 else entry.get('rate', 0)
            
            if name and rate:
                sq_yards = form_data['sq_yards']
                agreement_pct = form_data['agreement_percentage']
                
                total = rate * sq_yards
                agreement = total * agreement_pct
                registration = total - agreement
                
                c.execute("""
                    INSERT INTO commission_gm_entries 
                    (commission_id, name, total_amount, at_agreement, at_registration)
                    VALUES (%s, %s, %s, %s, %s)
                """, (commission_id, name.strip(), total, agreement, registration))
    
    # Save individual Agent entries

    
    conn.commit()
    conn.close()
    
    return commission_id


def update_commission_in_db(commission_id, form_data, calculations):
    """Update existing commission calculation in database"""
    conn = database.get_db_connection()
    c = conn.cursor()
    
    c.execute("""
        UPDATE commissions SET
            plot_no = %s, project_name = %s, sq_yards = %s, original_price = %s, negotiated_price = %s,
            advance_received = %s,
            agreement_percentage = %s,
            amount_paid_at_agreement = %s,
            amc_charges = %s,
            mediator_deduction = %s,
            broker_commission = %s,
            cgm_rate = %s,
            srgm_rate = %s,
            gm_rate = %s, dgm_rate = %s, agm_rate = %s,
            cgm_name = %s, srgm_name = %s, gm_name = %s, dgm_name = %s, agm_name = %s,
            total_amount = %s, w_value = %s, b_value = %s,
            balance_amount = %s, actual_agreement_amount = %s, agreement_balance = %s,
            cgm_total = %s, cgm_at_agreement = %s,
            cgm_at_registration = %s, srgm_total = %s, srgm_at_agreement = %s,
            srgm_at_registration = %s, gm_total = %s, gm_at_agreement = %s,
            gm_at_registration = %s, dgm_total = %s, dgm_at_agreement = %s, dgm_at_registration = %s,
            agm_total = %s, agm_at_agreement = %s, agm_at_registration = %s,
            commission_breakdown = %s
        WHERE id = %s
    """, (
        form_data['plot_no'], form_data.get('project_name', ''), form_data['sq_yards'], form_data['original_price'],
        form_data['negotiated_price'], form_data['advance_received'],
        form_data['agreement_percentage'],
        form_data['amount_paid_at_agreement'], form_data['amc_charges'], form_data.get('mediator_deduction', 0),
        form_data.get('broker_commission', 0),
        form_data['cgm_rate'], form_data['srgm_rate'], form_data['gm_rate'],
        form_data['dgm_rate'], form_data['agm_rate'], 
        form_data.get('cgm_name', ''), form_data.get('srgm_name', ''),
        form_data.get('gm_name', ''), form_data.get('dgm_name', ''), form_data.get('agm_name', ''),
        calculations['total_amount'], calculations['w_value'],
        calculations['b_value'], calculations['balance_amount'],
        calculations['actual_agreement_amount'], calculations['agreement_balance'],
        calculations['cgm_total'], calculations['cgm_at_agreement'],
        calculations['cgm_at_registration'], calculations['srgm_total'],
        calculations['srgm_at_agreement'], calculations['srgm_at_registration'],
        calculations['gm_total'], calculations['gm_at_agreement'],
        calculations['gm_at_registration'], 
        calculations['dgm_total'], calculations['dgm_at_agreement'], calculations['dgm_at_registration'],
        calculations['agm_total'], calculations['agm_at_agreement'], calculations['agm_at_registration'],
        form_data.get('commission_breakdown', ''),
        commission_id
    ))
    
    # Update DGM entries (Delete all and re-insert)
    c.execute("DELETE FROM commission_dgm_entries WHERE commission_id = %s", (commission_id,))
    dgm_entries = form_data.get('dgm_entries', [])
    if dgm_entries:
        for entry in dgm_entries:
            name = entry[0] if isinstance(entry, (list, tuple)) else entry.get('name', '')
            rate = entry[1] if isinstance(entry, (list, tuple)) and len(entry) > 1 else entry.get('rate', 0)
            
            if name and rate:
                sq_yards = form_data['sq_yards']
                agreement_pct = form_data['agreement_percentage']
                
                total = rate * sq_yards
                agreement = total * agreement_pct
                registration = total - agreement
                
                c.execute("""
                    INSERT INTO commission_dgm_entries 
                    (commission_id, name, total_amount, at_agreement, at_registration)
                    VALUES (%s, %s, %s, %s, %s)
                """, (commission_id, name.strip(), total, agreement, registration))

    # Update AGM entries (Delete all and re-insert)
    c.execute("DELETE FROM commission_agm_entries WHERE commission_id = %s", (commission_id,))
    agm_entries = form_data.get('agm_entries', [])
    if agm_entries:
        for entry in agm_entries:
            name = entry[0] if isinstance(entry, (list, tuple)) else entry.get('name', '')
            rate = entry[1] if isinstance(entry, (list, tuple)) and len(entry) > 1 else entry.get('rate', 0)
            
            if name and rate:
                sq_yards = form_data['sq_yards']
                agreement_pct = form_data['agreement_percentage']
                
                total = rate * sq_yards
                agreement = total * agreement_pct
                registration = total - agreement
                
                c.execute("""
                    INSERT INTO commission_agm_entries 
                    (commission_id, name, total_amount, at_agreement, at_registration)
                    VALUES (%s, %s, %s, %s, %s)
                """, (commission_id, name.strip(), total, agreement, registration))
    
    conn.commit()
    conn.close()



def generate_commission_pdf_bytes(form_data, calculations):
    """Generate commission PDF using ReportLab"""
    from reportlab.lib.pagesizes import A4
    from reportlab.lib import colors
    from reportlab.lib.units import inch
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.enums import TA_CENTER, TA_LEFT
    
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4, topMargin=0.5*inch, bottomMargin=0.5*inch)
    
    elements = []
    styles = getSampleStyleSheet()
    
    # Title
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=18,
        textColor=colors.HexColor('#f16924'),
        spaceAfter=12,
        alignment=TA_CENTER,
        fontName='Helvetica-Bold'
    )
    elements.append(Paragraph("Commission Calculation Report", title_style))
    elements.append(Spacer(1, 0.2*inch))
    
    # Plot Information
    plot_info = [
        ['Plot No:', form_data['plot_no']],
        ['Square Yards:', f"{form_data['sq_yards']:.2f}"],
        ['Original Price/Sq.Yd:', format_currency(form_data['original_price'])],
        ['Negotiated Price/Sq.Yd:', format_currency(form_data['negotiated_price'])],
    ]
    
    plot_table = Table(plot_info, colWidths=[2.5*inch, 3*inch])
    plot_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (0, -1), colors.HexColor('#f3f4f6')),
        ('TEXTCOLOR', (0, 0), (-1, -1), colors.black),
        ('ALIGN', (0, 0), (0, -1), 'LEFT'),
        ('ALIGN', (1, 0), (1, -1), 'RIGHT'),
        ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, -1), 10),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ('PADDING', (0, 0), (-1, -1), 8),
    ]))
    elements.append(plot_table)
    elements.append(Spacer(1, 0.3*inch))
    
    # Financial Summary
    subtitle_style = ParagraphStyle(
        'Subtitle',
        parent=styles['Heading2'],
        fontSize=14,
        textColor=colors.HexColor('#1f2937'),
        spaceAfter=6,
        fontName='Helvetica-Bold'
    )
    elements.append(Paragraph("Financial Summary", subtitle_style))
    elements.append(Spacer(1, 0.1*inch))
    
    financial_data = [
        ['Description', 'Amount'],
        ['Total Amount', format_currency(calculations['total_amount'])],
        ['W Value', format_currency(calculations['w_value'])],
        ['B Value', format_currency(calculations['b_value'])],
        ['Advance Received', format_currency(form_data['advance_received'])],
        ['Balance Amount', format_currency(calculations['balance_amount'])],
        ['Actual Agreement Amount', format_currency(calculations['actual_agreement_amount'])],
        ['Agreement Balance', format_currency(calculations['agreement_balance'])],
    ]
    
    financial_table = Table(financial_data, colWidths=[3*inch, 2.5*inch])
    financial_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#f16924')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (0, -1), 'LEFT'),
        ('ALIGN', (1, 0), (1, -1), 'RIGHT'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, 0), 11),
        ('FONTSIZE', (0, 1), (-1, -1), 10),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ('PADDING', (0, 0), (-1, -1), 8),
        ('BACKGROUND', (0, 1), (-1, -1), colors.white),
    ]))
    elements.append(financial_table)
    elements.append(Spacer(1, 0.3*inch))
    
    # Mediator Commission Summary Table
    elements.append(Paragraph("Mediator Commission", subtitle_style))
    elements.append(Spacer(1, 0.1*inch))
    
    # Initialize total calculations for distribution table
    # NEW LOGIC: Mediator Amount = Sq Yards * Broker Commission
    sq_yards = form_data['sq_yards']
    broker_commission = float(form_data.get('broker_commission') or 0)
    total_mediator_amount = broker_commission * sq_yards
    
    # NOTE: The loop below is kept for Distribution Table breakdown logic but 
    # it NO LONGER sums up to total_mediator_amount for the Summary Table.
    # The Summary Table uses the formula above.
    
    # We execute this loop just to prepare data structure if needed, or we rely on 'entries' being correct.
    # Actually, for the PDF summary 'total_mediator_amount' is the key.
    
    # ... Original loops kept for distribution table rendering ...
    
    distribution_total_check = 0 # To verify if needed, or just ignore.

    # CGM calc
    cgm_name = form_data.get('cgm_name', '') or '-'
    # sq_yards already defined above
    agreement_percentage = form_data['agreement_percentage']
    cgm_total = form_data['cgm_rate'] * sq_yards
    if form_data['cgm_rate'] > 0:
        pass # total_mediator_amount is already calculated via broker_commission
        
    srgm_entries = form_data.get('srgm_entries', [])
    for _, rate in srgm_entries:
        pass # total_mediator_amount is already calculated via broker_commission
        
    gm_entries = form_data.get('gm_entries', [])
    for _, rate in gm_entries:
        pass # total_mediator_amount is already calculated via broker_commission
        
    dgm_entries = form_data.get('dgm_entries', [])
    for _, rate in dgm_entries:
        pass # total_mediator_amount is already calculated via broker_commission
        
    agm_entries = form_data.get('agm_entries', [])
    for _, rate in agm_entries:
        pass # total_mediator_amount is already calculated via broker_commission

    deduction = form_data.get('mediator_deduction') or 0
    
    # Auto-calculate deduction if it's 0 (handling historic data)
    if deduction == 0:
        try:
            negotiated_price = float(form_data.get('negotiated_price') or 0)
            original_price = float(form_data.get('original_price') or 0)
            amc_charges = float(form_data.get('amc_charges') or 0)
            # Formula: (Original Price * Sq Yards) - ((Negotiated Price * Sq Yards) + (AMC Charges * Sq Yards))
            deduction = (original_price * sq_yards) - ((negotiated_price * sq_yards) + (amc_charges * sq_yards))
        except (ValueError, TypeError):
            deduction = 0

    actual_payment = total_mediator_amount - deduction
    at_agreement_payment = actual_payment * agreement_percentage

    mediator_data = [
        ['Description', 'Amount'],
        ['Mediator Amount', format_currency(total_mediator_amount)],
        ['Mediator Deduction', format_currency(deduction)],
        ['Actual Payment to Mediator', format_currency(actual_payment)],
        ['At Agreement', format_currency(at_agreement_payment)],
    ]
    
    mediator_table = Table(mediator_data, colWidths=[3*inch, 2.5*inch])
    mediator_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#f16924')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (0, -1), 'LEFT'),
        ('ALIGN', (1, 0), (1, -1), 'RIGHT'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, 0), 11),
        ('FONTSIZE', (0, 1), (-1, -1), 10),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ('PADDING', (0, 0), (-1, -1), 8),
        ('BACKGROUND', (0, 1), (-1, -1), colors.white),
    ]))
    elements.append(mediator_table)
    elements.append(Spacer(1, 0.3*inch))

    # Commission Distribution Table (Detailed Breakdown)
    elements.append(Paragraph("Commission Distribution", subtitle_style))
    elements.append(Spacer(1, 0.1*inch))
    
    # Build distribution data with individual rows
    distribution_data = [
        ['Name', 'Role', 'Total', 'At Agreement', 'At Registration'],
    ]
    
    # Re-using variables from top
    
    cgm_agreement = cgm_total * agreement_percentage
    cgm_registration = cgm_total - cgm_agreement
    
    if form_data['cgm_rate'] > 0:
        distribution_data.append([
            cgm_name, 'CGM', 
            format_currency(cgm_total),
            format_currency(cgm_agreement),
            format_currency(cgm_registration)
        ])
    
    # Add individual Sr. GM rows
    for name, rate in srgm_entries:
        total = rate * sq_yards
        agreement = total * agreement_percentage
        registration = total - agreement
        distribution_data.append([
            name or '-', 'SrGM',
            format_currency(total),
            format_currency(agreement),
            format_currency(registration)
        ])
    
    # Add individual GM rows
    for name, rate in gm_entries:
        total = rate * sq_yards
        agreement = total * agreement_percentage
        registration = total - agreement
        distribution_data.append([
            name or '-', 'GM',
            format_currency(total),
            format_currency(agreement),
            format_currency(registration)
        ])
    
    # Add individual DGM rows
    for name, rate in dgm_entries:
        total = rate * sq_yards
        agreement = total * agreement_percentage
        registration = total - agreement
        distribution_data.append([
            name or '-', 'DGM',
            format_currency(total),
            format_currency(agreement),
            format_currency(registration)
        ])
    
    # Add individual AGM rows
    for name, rate in agm_entries:
        total = rate * sq_yards
        agreement = total * agreement_percentage
        registration = total - agreement
        distribution_data.append([
            name or '-', 'AGM',
            format_currency(total),
            format_currency(agreement),
            format_currency(registration)
        ])
    
    distribution_table = Table(distribution_data, colWidths=[1.3*inch, 1.0*inch, 1.2*inch, 1.2*inch, 1.3*inch])
    distribution_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#10b981')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (1, -1), 'LEFT'),  # Name and Role columns left-aligned
        ('ALIGN', (2, 0), (-1, -1), 'RIGHT'),  # Numeric columns right-aligned
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTNAME', (0, 1), (1, -1), 'Helvetica-Bold'),  # Name and Role bold
        ('FONTSIZE', (0, 0), (-1, 0), 10),
        ('FONTSIZE', (0, 1), (-1, -1), 9),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ('PADDING', (0, 0), (-1, -1), 8),
        ('BACKGROUND', (0, 1), (-1, -1), colors.white),
    ]))
    elements.append(distribution_table)
    
    # Build PDF
    doc.build(elements)
    
    pdf_bytes = buffer.getvalue()
    buffer.close()
    
    return pdf_bytes


def generate_commission_docx_bytes(form_data, calculations):
    """Generate commission Word document using python-docx"""
    doc = docx.Document()
    
    # Title
    title = doc.add_heading('Commission Calculation Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Plot Information
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Plot Details'
    hdr_cells[1].text = 'Values'
    
    plot_info = [
        ('Plot No:', form_data['plot_no']),
        ('Square Yards:', f"{form_data['sq_yards']:.2f}"),
        ('Original Price/Sq.Yd:', format_currency(form_data['original_price'])),
        ('Negotiated Price/Sq.Yd:', format_currency(form_data['negotiated_price'])),
    ]
    
    for label, value in plot_info:
        row_cells = table.add_row().cells
        row_cells[0].text = label
        row_cells[1].text = str(value)
        
    doc.add_paragraph()
    
    # Financial Summary
    doc.add_heading('Financial Summary', level=2)
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Description'
    hdr_cells[1].text = 'Amount'
    
    financial_data = [
        ('Total Amount', format_currency(calculations['total_amount'])),
        ('W Value', format_currency(calculations['w_value'])),
        ('B Value', format_currency(calculations['b_value'])),
        ('Balance Amount', format_currency(calculations['balance_amount'])),
        ('Actual Agreement Amount', format_currency(calculations['actual_agreement_amount'])),
        ('Agreement Balance', format_currency(calculations['agreement_balance'])),
    ]
    
    for label, value in financial_data:
        row_cells = table.add_row().cells
        row_cells[0].text = label
        row_cells[1].text = str(value)
        
    doc.add_paragraph()
    

    
    # Initialize total calculations for docx
    # NEW LOGIC: Mediator Amount = Sq Yards * Broker Commission
    sq_yards = form_data['sq_yards']
    broker_commission = float(form_data.get('broker_commission') or 0)
    total_mediator_amount = broker_commission * sq_yards
    
    # Calculate totals first for summary - loops below no longer sum to total_mediator_amount
    
    # CGM calc
    cgm_name = form_data.get('cgm_name', '') or '-'
    # sq_yards already defined above
    agreement_percentage = form_data['agreement_percentage']
    cgm_total = form_data['cgm_rate'] * sq_yards
    if form_data['cgm_rate'] > 0:
        pass # total_mediator_amount is already calculated
        
    srgm_entries = form_data.get('srgm_entries', [])
    for _, rate in srgm_entries:
        pass # total_mediator_amount is already calculated
        
    gm_entries = form_data.get('gm_entries', [])
    for _, rate in gm_entries:
        pass # total_mediator_amount is already calculated
        
    dgm_entries = form_data.get('dgm_entries', [])
    for _, rate in dgm_entries:
        pass # total_mediator_amount is already calculated
        
    agm_entries = form_data.get('agm_entries', [])
    for _, rate in agm_entries:
        pass # total_mediator_amount is already calculated

    # Mediator Commission Summary Table (Moved to Top)
    doc.add_heading('Mediator Commission', level=2)
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Description'
    hdr_cells[1].text = 'Amount'
    
    deduction = form_data.get('mediator_deduction') or 0
    
    # Auto-calculate deduction if it's 0 (handling historic data)
    if deduction == 0:
        try:
            negotiated_price = float(form_data.get('negotiated_price') or 0)
            original_price = float(form_data.get('original_price') or 0)
            amc_charges = float(form_data.get('amc_charges') or 0)
            # Formula: (Original Price * Sq Yards) - ((Negotiated Price * Sq Yards) + (AMC Charges * Sq Yards))
            deduction = (original_price * sq_yards) - ((negotiated_price * sq_yards) + (amc_charges * sq_yards))
        except (ValueError, TypeError):
            deduction = 0
            
    actual_payment = total_mediator_amount - deduction
    at_agreement_payment = actual_payment * agreement_percentage
    
    mediator_data = [
        ('Mediator Amount', format_currency(total_mediator_amount)),
        ('Mediator Deduction', format_currency(deduction)),
        ('Actual Payment to Mediator', format_currency(actual_payment)),
        ('At Agreement', format_currency(at_agreement_payment)),
    ]
    
    for label, value in mediator_data:
        row_cells = table.add_row().cells
        row_cells[0].text = label
        row_cells[1].text = str(value)
        
    doc.add_paragraph()

    # Commission Distribution Table (Detailed Breakdown)
    doc.add_heading('Commission Distribution', level=2)
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Name'
    hdr_cells[1].text = 'Role'
    hdr_cells[2].text = 'Total'
    hdr_cells[3].text = 'At Agreement'
    hdr_cells[4].text = 'At Registration'
    
    # Re-using variables from top
    
    cgm_agreement = cgm_total * agreement_percentage
    cgm_registration = cgm_total - cgm_agreement
    
    if form_data['cgm_rate'] > 0:
        row_cells = table.add_row().cells
        row_cells[0].text = cgm_name
        row_cells[1].text = 'CGM'
        row_cells[2].text = format_currency(cgm_total)
        row_cells[3].text = format_currency(cgm_agreement)
        row_cells[4].text = format_currency(cgm_registration)
    
    # Add individual Sr. GM rows
    for name, rate in srgm_entries:
        total = rate * sq_yards
        agreement = total * agreement_percentage
        registration = total - agreement
        row_cells = table.add_row().cells
        row_cells[0].text = name or '-'
        row_cells[1].text = 'SrGM'
        row_cells[2].text = format_currency(total)
        row_cells[3].text = format_currency(agreement)
        row_cells[4].text = format_currency(registration)
    
    # Add individual GM rows
    for name, rate in gm_entries:
        total = rate * sq_yards
        agreement = total * agreement_percentage
        registration = total - agreement
        row_cells = table.add_row().cells
        row_cells[0].text = name or '-'
        row_cells[1].text = 'GM'
        row_cells[2].text = format_currency(total)
        row_cells[3].text = format_currency(agreement)
        row_cells[4].text = format_currency(registration)
    
    # Add individual DGM rows
    for name, rate in dgm_entries:
        total = rate * sq_yards
        agreement = total * agreement_percentage
        registration = total - agreement
        row_cells = table.add_row().cells
        row_cells[0].text = name or '-'
        row_cells[1].text = 'DGM'
        row_cells[2].text = format_currency(total)
        row_cells[3].text = format_currency(agreement)
        row_cells[4].text = format_currency(registration)
    
    # Add individual AGM rows
    for name, rate in agm_entries:
        total = rate * sq_yards
        agreement = total * agreement_percentage
        registration = total - agreement
        row_cells = table.add_row().cells
        row_cells[0].text = name or '-'
        row_cells[1].text = 'AGM'
        row_cells[2].text = format_currency(total)
        row_cells[3].text = format_currency(agreement)
        row_cells[4].text = format_currency(registration)
        
          
    # Save to bytes
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()


@app.route("/commission/download_docx/<int:commission_id>/<filename>")
def download_commission_docx(commission_id, filename):
    """Download commission as Word document"""
    conn = database.get_db_connection()
    c = conn.cursor()
    
    c.execute("SELECT * FROM commissions WHERE id = %s", (commission_id,))
    row = database.fetch_one(c)
    conn.close()
    
    if not row:
        abort(404)
        
    # Reconstruct data
    form_data = {
        'plot_no': row['plot_no'],
        'sq_yards': row['sq_yards'],
        'original_price': row['original_price'],
        'negotiated_price': row['negotiated_price'],
        'cgm_name': row.get('cgm_name', ''),
        'srgm_name': row.get('srgm_name', ''),
        'gm_name': row.get('gm_name', ''),
        'agent_name': row.get('agent_name', ''),
    }
    
    calculations = {
        'total_amount': row['total_amount'],
        'w_value': row['w_value'],
        'b_value': row['b_value'],
        'balance_amount': row['balance_amount'],
        'actual_agreement_amount': row['actual_agreement_amount'],
        'agreement_balance': row['agreement_balance'],
        'mediator_amount': row['mediator_amount'],
        'mediator_deduction': row['mediator_deduction'],
        'mediator_actual_payment': row['mediator_actual_payment'],
        'mediator_at_agreement': row['mediator_at_agreement'],
        'cgm_total': row['cgm_total'],
        'cgm_at_agreement': row['cgm_at_agreement'],
        'cgm_at_registration': row['cgm_at_registration'],
        'srgm_total': row['srgm_total'],
        'srgm_at_agreement': row['srgm_at_agreement'],
        'srgm_at_registration': row['srgm_at_registration'],
        'gm_total': row['gm_total'],
        'gm_at_agreement': row['gm_at_agreement'],
        'gm_at_registration': row['gm_at_registration'],
        'agent_total': row['agent_total'],
        'agent_at_agreement': row['agent_at_agreement'],
        'agent_at_registration': row['agent_at_registration'],
    }
    
    docx_bytes = generate_commission_docx_bytes(form_data, calculations)
    
    plot_no_safe = _safe_filename_fragment(form_data['plot_no'], 'commission')
    filename = f"commission_Plot{plot_no_safe}.docx"
    
    response = make_response(docx_bytes)
    response.headers['Content-Type'] = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    response.headers['Content-Disposition'] = f'attachment; filename="{filename}"'
    
    return response
    return response


@app.route("/commission/search", methods=["GET", "POST"])
def search_commission():
    """Search for commission by plot number and project"""
    projects = get_projects()
    
    if request.method == "POST":
        plot_no = request.form.get("plot_no", "").strip()
        project_name = request.form.get("project_name", "").strip()
        
        conn = database.get_db_connection()
        c = conn.cursor()
        
        # Find latest commission for this plot and project
        if project_name:
            c.execute("""
                SELECT id FROM commissions 
                WHERE plot_no = %s AND project_name = %s
                ORDER BY id DESC LIMIT 1
            """, (plot_no, project_name))
        else:
            c.execute("""
                SELECT id FROM commissions 
                WHERE plot_no = %s 
                ORDER BY id DESC LIMIT 1
            """, (plot_no,))
        
        row = database.fetch_one(c)
        conn.close()
        
        if row:
            return redirect(url_for('preview_commission_pdf', commission_id=row[0]))
        else:
            if project_name:
                flash(f"No commission calculation found for Plot {plot_no} in {project_name}", "warning")
            else:
                flash(f"No commission calculation found for Plot {plot_no}", "warning")
            return redirect(url_for('search_commission'))
            
    return render_template("commission_search.html", projects=projects)


@app.route("/commissions/view")
def view_commissions():
    """View commissions by project and plot"""
    if not (session.get("role") == "admin" or session.get("can_view_dashboard")):
        abort(403)
        
    projects = get_projects()
    selected_project = request.args.get("project")
    search_plot = request.args.get("search_plot", "").strip()
    plots = []
    
    # Always fetch plots, applying filters if they exist
    conn = database.get_db_connection()
    c = conn.cursor()
    
    # Build query based on filters
    where_clauses = []
    params = []
    
    if selected_project:
        where_clauses.append("project_name = %s")
        params.append(selected_project)
    
    if search_plot:
        where_clauses.append("plot_no LIKE %s")
        params.append(f"%{search_plot}%")
    
    where_sql = " AND ".join(where_clauses) if where_clauses else "1=1"
    
    # Get all plots that have commissions
    c.execute(f"""
        SELECT DISTINCT plot_no, project_name, cgm_name, srgm_name, gm_name, total_amount
        FROM commissions 
        WHERE {where_sql}
        AND project_name IS NOT NULL AND project_name != ''
        ORDER BY CAST(plot_no AS UNSIGNED), plot_no
    """, tuple(params))
    
    rows = database.fetch_all(c)
    conn.close()
    
    # Convert to list of dicts
    for r in rows:
        plots.append({
            'plot_no': r['plot_no'],
            'project_name': r['project_name'],
            'cgm_name': r['cgm_name'],
            'srgm_name': r['srgm_name'],
            'gm_name': r['gm_name'],

            'total_commission': r['total_amount']
        })
            
    return render_template("commission_view.html", 
                         projects=projects, 
                         selected_project=selected_project,
                         search_plot=search_plot,
                         plots=plots)


@app.route("/commissions/view/<project_name>/<plot_no>")
def view_commission_detail(project_name, plot_no):
    """View commission details for a specific plot"""
    if not (session.get("role") == "admin" or session.get("can_view_dashboard")):
        abort(403)
        
    conn = database.get_db_connection()
    c = conn.cursor()
    
    # Get commission details
    c.execute("""
        SELECT * FROM commissions 
        WHERE project_name = %s AND plot_no = %s
        ORDER BY id DESC LIMIT 1
    """, (project_name, plot_no))
    
    commission = database.fetch_one(c)
    conn.close()
    
    if not commission:
        flash(f"No commission found for Plot {plot_no} in {project_name}", "warning")
        return redirect(url_for('view_commissions', project=project_name))
        
    # Reuse the preview template or create a new one
    # For now, let's redirect to the existing preview page which shows details nicely
    return redirect(url_for('preview_commission_pdf', commission_id=commission['id']))


@app.route("/plot-layout/<project_name>")
def plot_layout_viewer(project_name):
    """Interactive plot layout viewer with real-time status"""
    # Access control: Admin OR user with 'can_view_vishvam_layout' permission (for Vishvam project)
    if session.get("role") != "admin":
        if project_name.lower() == "vishvam" and not session.get("can_view_vishvam_layout"):
            abort(403)
        # For other projects, default to admin-only or dashboard view permission (adjust as needed)
        elif project_name.lower() != "vishvam" and not session.get("can_view_dashboard"):
             abort(403)
    
    # Get all plots with their status from receipts
    conn = database.get_db_connection()
    c = conn.cursor()
    
    # Get sold plots
    c.execute("""
        SELECT DISTINCT plot_no 
        FROM receipts 
        WHERE project_name = %s AND plot_no IS NOT NULL AND plot_no != ''
    """, (project_name,))
    
    rows = database.fetch_all(c)
    print(f"DEBUG: Raw rows from database: {rows}")
    
    sold_plots = []
    for row in rows:
        if row and 'plot_no' in row:
            plot_no = row['plot_no']
            if plot_no:  # Make sure it's not None or empty
                sold_plots.append(str(plot_no))
    
    print(f"DEBUG: Final sold_plots list: {sold_plots}")
    print(f"DEBUG: sold_plots type: {type(sold_plots)}")
    print(f"DEBUG: sold_plots length: {len(sold_plots)}")
    
    # Get plot metadata if exists
    c.execute("""
        SELECT plot_no, facing, length, width, area, sq_yards, status, notes, svg_element_id,
               boundary_east, boundary_west, boundary_north, boundary_south
        FROM plot_layouts
        WHERE project_name = %s
    """, (project_name,))
    plot_metadata = {row['plot_no']: dict(row) for row in database.fetch_all(c)}
    
    conn.close()
    
    # Check if SVG layout exists (preferred) or PDF fallback
    svg_path = f"layouts/{project_name.lower()}_layout.svg"
    pdf_path = f"layouts/{project_name.lower()}_layout.pdf"
    
    # Check if SVG exists first
    import os
    svg_full_path = os.path.join('static', svg_path)
    pdf_full_path = os.path.join('static', pdf_path)
    
    print(f"DEBUG: Checking for layout files...")
    print(f"DEBUG: SVG path: {svg_full_path}")
    print(f"DEBUG: PDF path: {pdf_full_path}")
    print(f"DEBUG: SVG exists: {os.path.exists(svg_full_path)}")
    print(f"DEBUG: PDF exists: {os.path.exists(pdf_full_path)}")
    print(f"DEBUG: Current working dir: {os.getcwd()}")
    
    if os.path.exists(svg_full_path):
        layout_file = url_for('static', filename=svg_path)
        layout_type = 'svg'
        print(f"Loading SVG layout: {layout_file}")
    elif os.path.exists(pdf_full_path):
        layout_file = url_for('static', filename=pdf_path)
        layout_type = 'pdf'
        print(f"Loading PDF layout: {layout_file}")
    else:
        # No layout file exists - show message
        print(f"No layout file found for project: {project_name}")
        return render_template("plot_layout_viewer.html",
                             project_name=project_name,
                             layout_file=None,
                             layout_type=None,
                             sold_plots=sold_plots,
                             plot_metadata=plot_metadata,
                             user_role=session.get("role", ""))
    
    print(f"DEBUG: Returning with layout_file={layout_file}, layout_type={layout_type}")
    return render_template("plot_layout_viewer.html",
                         project_name=project_name,
                         layout_file=layout_file,
                         layout_type=layout_type,
                         sold_plots=sold_plots,
                         plot_metadata=plot_metadata,
                         user_role=session.get("role", ""))


@app.route("/api/plot-status/<project_name>")
def get_plot_status(project_name):
    """API endpoint to get real-time plot status"""
    conn = database.get_db_connection()
    c = conn.cursor()
    
    # Get sold plots
    c.execute("""
        SELECT DISTINCT plot_no 
        FROM receipts 
        WHERE project_name = %s AND plot_no IS NOT NULL AND plot_no != ''
    """, (project_name,))
    sold_plots = [row['plot_no'] for row in database.fetch_all(c)]
    
    conn.close()
    
    return jsonify({
        'sold_plots': sold_plots,
        'timestamp': datetime.now().isoformat()
    })


@app.route("/api/plot-mapping/save", methods=["POST"])
def save_plot_mapping():
    """API endpoint to save plot mapping metadata (Admin only)"""
    # Only admins can edit plot mappings
    if not session.get("role") == "admin":
        return jsonify({'error': 'Unauthorized. Only admins can edit plot mappings.'}), 403
        
    data = request.json
    project_name = data.get('project_name')
    plot_no = data.get('plot_no')
    
    if not project_name or not plot_no:
        return jsonify({'error': 'Missing required fields'}), 400
        
    conn = database.get_db_connection()
    c = conn.cursor()
    
    # Get all fields
    facing = data.get('facing')
    length = float(data.get('length') or 0)
    width = float(data.get('width') or 0)
    sq_yards = float(data.get('sq_yards') or 0)
    status = data.get('status', 'available')  # Default to available
    boundary_east = data.get('boundary_east', '')
    boundary_west = data.get('boundary_west', '')
    boundary_north = data.get('boundary_north', '')
    boundary_south = data.get('boundary_south', '')
    
    # Calculate area if dimensions provided (fallback if sq_yards not provided)
    area = sq_yards if sq_yards > 0 else (length * width if length and width else 0)
    
    try:
        # Check if exists
        c.execute("SELECT id FROM plot_layouts WHERE project_name = %s AND plot_no = %s", (project_name, plot_no))
        existing = database.fetch_one(c)
        
        if existing:
            c.execute("""
                UPDATE plot_layouts 
                SET facing = %s, length = %s, width = %s, area = %s, sq_yards = %s, status = %s, 
                    boundary_east = %s, boundary_west = %s, boundary_north = %s, boundary_south = %s,
                    svg_element_id = %s
                WHERE project_name = %s AND plot_no = %s
            """, (facing, length, width, area, sq_yards, status, 
                  boundary_east, boundary_west, boundary_north, boundary_south,
                  data.get('element_id'), project_name, plot_no))
        else:
            c.execute("""
                INSERT INTO plot_layouts (project_name, plot_no, facing, length, width, area, sq_yards, status, 
                                         boundary_east, boundary_west, boundary_north, boundary_south, svg_element_id)
                VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
            """, (project_name, plot_no, facing, length, width, area, sq_yards, status,
                  boundary_east, boundary_west, boundary_north, boundary_south, data.get('element_id')))
            
        conn.commit()
        return jsonify({'success': True, 'message': f'Plot {plot_no} mapped successfully', 'status': status})
    except Exception as e:
        return jsonify({'error': str(e)}), 500
    finally:
        conn.close()


# -----------------------------
# Debug route
# -----------------------------
@app.route("/_routes")
def list_routes():
    rules = []
    for rule in app.url_map.iter_rules():
        rules.append({
            "rule": str(rule),
            "endpoint": rule.endpoint,
            "methods": sorted([m for m in rule.methods if m not in ("HEAD", "OPTIONS")])
        })
    return jsonify({"routes": rules})


# -----------------------------
# Start Server
# -----------------------------

# -----------------------------
# Create Receipts Options
# -----------------------------
@app.route("/create_receipts")
def create_receipts_options():
    return render_template("create_receipt_options.html")

# -----------------------------
# Import Receipts (Bulk Upload)
# -----------------------------
@app.route("/import_receipts", methods=["GET", "POST"])
def import_receipts():
    if not (session.get("role") == "admin"):
        flash("Access denied. Admin privileges required.", "danger")
        return redirect(url_for('dashboard'))

    if request.method == "POST":
        if 'file' not in request.files:
            flash('No file part', 'danger')
            return redirect(request.url)
            
        file = request.files['file']
        if file.filename == '':
            flash('No selected file', 'danger')
            return redirect(request.url)
            
        if file and (file.filename.endswith('.xlsx') or file.filename.endswith('.xls')):
            try:
                # Save file temporarily to process in next step
                filename = f"import_{int(datetime.now().timestamp())}_{file.filename}"
                filepath = os.path.join(app.config["UPLOAD_FOLDER"], filename)
                file.save(filepath)
                
                # Load headers
                wb = openpyxl.load_workbook(filepath, data_only=True)
                ws = wb.active
                # Assume headers are in Row 2 (Index 1) as per standard project file, 
                # or Row 1 if standard excel. Let's grab first few rows and try to find header row?
                # For simplicity/consistency with previous analysis: Reading Row 2 (index 1) as headers.
                # But to be safer, we can pass the first few rows to UI? 
                # Let's just grab Row 2 as headers for now, matching the standard format.
                rows = list(ws.iter_rows(values_only=True))
                
                headers = []
                if len(rows) > 1:
                    headers = list(rows[1]) # Row 2
                elif len(rows) > 0:
                    headers = list(rows[0]) # Row 1 if small file
                    
                # Clean headers (None -> Column X)
                clean_headers = []
                for idx, h in enumerate(headers):
                    clean_headers.append(str(h) if h else f"Column {idx+1}")
                
                return render_template("import_mapping.html", headers=clean_headers, filename=filename)
                
            except Exception as e:
                flash(f"Error reading file: {str(e)}", "danger")
                return redirect(request.url)

    return render_template("import_receipts.html")


@app.route("/import_receipts/process", methods=["POST"])
def process_import_mapping():
    if not (session.get("role") == "admin"):
        abort(403)
        
    filename = request.form.get("filename")
    filepath = os.path.join(app.config["UPLOAD_FOLDER"], filename)
    
    if not os.path.exists(filepath):
        flash("File processing error: Temporary file not found.", "danger")
        return redirect(url_for("import_receipts"))

    # Get mappings (Column Indices)
    try:
        col_plot = int(request.form.get("col_plot_no"))
        col_date = int(request.form.get("col_date"))
        col_amount = int(request.form.get("col_amount"))
        
        # Optional fields
        col_customer = request.form.get("col_customer")
        col_payment_mode = request.form.get("col_payment_mode")
        col_receipt_no = request.form.get("col_receipt_no")
        col_sq_yards = request.form.get("col_sq_yards")
        
        col_customer = int(col_customer) if col_customer else None
        col_payment_mode = int(col_payment_mode) if col_payment_mode else None
        col_receipt_no = int(col_receipt_no) if col_receipt_no else None
        col_sq_yards = int(col_sq_yards) if col_sq_yards else None
        
        col_basic_price = request.form.get("col_basic_price")
        col_basic_price = int(col_basic_price) if col_basic_price else None
        
        import_summary = {
            'total_rows': 0,
            'imported': 0,
            'skipped': 0,
            'errors': []
        }
        
        wb = openpyxl.load_workbook(filepath, data_only=True)
        ws = wb.active
        rows = list(ws.iter_rows(values_only=True))
        
        # Context for grouped data
        current_plot_no = None
        current_sq_yards = None
        current_basic_price = None
        current_customer_name = None
        
        conn = database.get_db_connection()
        c = conn.cursor()
        
        # Start iterating from Row 3 (Index 2) - Data
        for i, row in enumerate(rows[2:], start=3):
            import_summary['total_rows'] += 1
            if not row: continue

            # 1. Plot No
            if col_plot < len(row):
                raw_plot = row[col_plot]
                if raw_plot not in [None, 0, '0', 0.0, '']:
                    current_plot_no = str(raw_plot).strip()
                    if current_plot_no.endswith('.0'):
                        current_plot_no = current_plot_no[:-2]
                    # Reset context specific to new plot
                    current_customer_name = None 
                    current_sq_yards = None
                    current_basic_price = None
                    
                    # Update customer if mapped and present on this plot row
                    if col_customer is not None and col_customer < len(row) and row[col_customer]:
                         current_customer_name = str(row[col_customer]).strip()
                                        # Update sq yards if mapped and present
                    if col_sq_yards is not None and col_sq_yards < len(row) and row[col_sq_yards]:
                         current_sq_yards = str(row[col_sq_yards]).strip()

                    # Update basic price if mapped and present
                    if col_basic_price is not None and col_basic_price < len(row) and row[col_basic_price]:
                         try:
                             current_basic_price = float(row[col_basic_price])
                         except (ValueError, TypeError):
                             pass


            # 2. Amount
            amount = 0
            if col_amount < len(row):
                raw_amount = row[col_amount]
                try:
                    amount = float(raw_amount) if raw_amount else 0
                except (ValueError, TypeError):
                    amount = 0
            
            if amount > 0:
                if not current_plot_no:
                    import_summary['errors'].append(f"Row {i}: Found amount {amount} but no Plot No context.")
                    continue
                
                # 3. Date
                date_str = datetime.now().strftime("%Y-%m-%d")
                if col_date < len(row):
                    raw_date = row[col_date]
                    if isinstance(raw_date, datetime):
                        date_str = raw_date.strftime("%Y-%m-%d")
                    elif isinstance(raw_date, str):
                        try:
                            # Try common formats
                            for fmt in ["%d.%m.%Y", "%Y-%m-%d", "%d-%m-%Y", "%m/%d/%Y"]:
                                try:
                                    date_str = datetime.strptime(raw_date, fmt).strftime("%Y-%m-%d")
                                    break
                                except ValueError:
                                    continue
                        except:
                            pass
                            
                # 4. Receipt No
                receipt_no = None
                if col_receipt_no is not None and col_receipt_no < len(row):
                    raw_no = row[col_receipt_no]
                    if raw_no:
                        # integer cleaning
                        digits = re.findall(r'\d+', str(raw_no))
                        if digits:
                            receipt_no = "".join(digits)
                        else:
                            receipt_no = str(raw_no).strip()
                            
                if not receipt_no:
                     receipt_no = f"IMP-{current_plot_no}-{int(datetime.now().timestamp())}-{i}"

                # 5. Payment Mode
                payment_mode = "Unknown"
                if col_payment_mode is not None and col_payment_mode < len(row):
                    val = row[col_payment_mode]
                    if val:
                        payment_mode = str(val).strip()

                # Deduplication
                exists = False
                c.execute("SELECT id FROM receipts WHERE no = %s OR instrument_no = %s", (receipt_no, receipt_no))
                if database.fetch_one(c):
                    exists = True
                
                if exists:
                    import_summary['skipped'] += 1
                else:
                    try:
                        c.execute("""
                            INSERT INTO receipts (
                                date, plot_no, customer_name, amount_numeric, amount_words, 
                                payment_mode, no, project_name, square_yards, basic_price
                            ) VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
                        """, (
                            date_str, 
                            current_plot_no, 
                            current_customer_name or "Unknown", 
                            amount, 
                            number_to_words(amount), 
                            payment_mode, 
                            receipt_no, 
                            "Vishvam", 
                            current_sq_yards,
                            current_basic_price
                        ))
                        import_summary['imported'] += 1
                    except Exception as e:
                        import_summary['errors'].append(f"Row {i}: DB Error - {str(e)}")

        conn.commit()
        conn.close()
        
        # cleanup
        try:
            os.remove(filepath)
        except:
            pass
            
        return render_template("import_receipts.html", summary=import_summary)

    except Exception as e:
        flash(f"Error processing import: {str(e)}", "danger")
        return redirect(url_for("import_receipts"))

@app.route("/delete_receipts")
def delete_receipts():
    """List plots for receipt deletion"""
    if not (session.get("role") == "admin"):
        flash("Access denied. Admin privileges required.", "danger")
        return redirect(url_for('dashboard'))
        
    conn = database.get_db_connection()
    c = conn.cursor()
    
    # Get plots with receipts count
    c.execute("""
        SELECT plot_no, project_name, COUNT(*) as receipt_count 
        FROM receipts 
        WHERE plot_no IS NOT NULL AND plot_no != '' 
        GROUP BY plot_no, project_name 
        ORDER BY CAST(plot_no AS UNSIGNED)
    """)
    
    rows = database.fetch_all(c)
    conn.close()
    
    plots = []
    for r in rows:
        plots.append({
            'plot_no': r['plot_no'],
            'project_name': r.get('project_name', ''),
            'receipt_count': r['receipt_count']
        })
        
    return render_template("delete_receipts_list.html", plots=plots)


@app.route("/delete_receipts/<plot_no>")
def delete_receipts_detail(plot_no):
    """View receipts for deletion for a specific plot"""
    if not (session.get("role") == "admin"):
        flash("Access denied. Admin privileges required.", "danger")
        return redirect(url_for('dashboard'))
        
    project_name = request.args.get("project_name", "")
    
    conn = database.get_db_connection()
    c = conn.cursor()
    
    # Get receipts
    if project_name:
        c.execute("""
            SELECT * FROM receipts 
            WHERE plot_no = %s AND project_name = %s
            ORDER BY date DESC, id DESC
        """, (plot_no, project_name))
    else:
        c.execute("""
            SELECT * FROM receipts 
            WHERE plot_no = %s 
            ORDER BY date DESC, id DESC
        """, (plot_no,))
        
    rows = database.fetch_all(c)
    conn.close()
    
    receipts = []
    for row in rows:
        r = dict_from_row(row)
        r["amount_formatted"] = format_inr(r["amount_numeric"])
        receipts.append(r)
        
    return render_template("delete_receipts_detail.html", 
                         receipts=receipts, 
                         plot_no=plot_no, 
                         project_name=project_name)


@app.route("/receipts/bulk_delete", methods=["POST"])
def bulk_delete_receipts():
    """Handle bulk deletion of receipts"""
    if not (session.get("role") == "admin"):
        flash("Access denied. Admin privileges required.", "danger")
        return redirect(url_for('dashboard'))
        
    receipt_ids = request.form.getlist("receipt_ids")
    project_name = request.form.get("project_name")
    
    if not receipt_ids:
        flash("No receipts selected for deletion", "warning")
        return redirect(request.referrer)
        
    try:
        conn = database.get_db_connection()
        c = conn.cursor()
        
        # Convert IDs to integers for safety
        ids = [int(rid) for rid in receipt_ids]
        format_strings = ','.join(['%s'] * len(ids))
        
        # Delete
        c.execute(f"DELETE FROM receipts WHERE id IN ({format_strings})", tuple(ids))
        deleted_count = c.rowcount
        conn.commit()
        conn.close()
        
        flash(f"Successfully deleted {deleted_count} receipt(s).", "success")
        
    except Exception as e:
        flash(f"Error deleting receipts: {str(e)}", "danger")
        app.logger.error(f"Bulk delete error: {e}")
        
    return redirect(request.referrer)

@app.route("/receipts/bulk_delete_plots", methods=["POST"])
def bulk_delete_plots():
    """Handle bulk deletion of all receipts for selected plots"""
    if not (session.get("role") == "admin"):
        flash("Access denied. Admin privileges required.", "danger")
        return redirect(url_for('dashboard'))
        
    plot_data_list = request.form.getlist("plot_data")
    
    if not plot_data_list:
        flash("No plots selected for deletion", "warning")
        return redirect(request.referrer)
        
    try:
        conn = database.get_db_connection()
        c = conn.cursor()
        
        deleted_count = 0
        plots_affected = 0
        
        for item in plot_data_list:
            # Parse "plot_no|project_name" or just "plot_no|"
            parts = item.split('|')
            plot_no = parts[0]
            project_name = parts[1] if len(parts) > 1 else ""
            
            if project_name:
                c.execute("DELETE FROM receipts WHERE plot_no = %s AND project_name = %s", (plot_no, project_name))
            else:
                c.execute("DELETE FROM receipts WHERE plot_no = %s", (plot_no,))
                
            deleted_count += c.rowcount
            plots_affected += 1
            
        conn.commit()
        conn.close()
        
        flash(f"Successfully deleted {deleted_count} receipts across {plots_affected} plots.", "success")
        
    except Exception as e:
        flash(f"Error deleting plots: {str(e)}", "danger")
        app.logger.error(f"Bulk plot delete error: {e}")
        
    return redirect(request.referrer)

if __name__ == "__main__":
    init_db()
    os.makedirs(app.config["UPLOAD_FOLDER"], exist_ok=True)
    os.makedirs(os.path.join("static", "css"), exist_ok=True)
    app.run(debug=True, port=5000)

